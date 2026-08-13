"""Auditor-input endpoints for the MANUAL audit workspace.

These power the Manual copy of the Learner Log Pro app
(``/workspace/auditor-manual``). The activity feed itself is served by
``ledger_views`` from the ``Manual_audit`` schema; this module owns the
auditor-entered data:

* per-activity annotations (planned-hours override + KSB notes)
  -> ``"Manual_audit".activity_annotations``
* the create/replace/delete activity overlay
  -> ``"Manual_audit".activity_overrides``
* the rich learner profile (contracts, evidence, ILR, skills radar, ...)
  reading the shared reference sources (``fetching_evidence``,
  ``Audit.learner_match.aptem_training_plan``, ``Audit.ilr_learning_deliveries``)
  but resolving learners from ``"Manual_audit".learners`` and storing every
  write in ``Manual_audit`` tables.

The module intentionally does not import from ``audit_api`` so the automatic
and manual audit systems stay independent.
"""

import datetime
import io
import json
import math
import mimetypes
import re
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from urllib.parse import quote, unquote, urlparse

from django.conf import settings
from django.db import DatabaseError, connections, transaction
from django.http import HttpRequest, JsonResponse

from audit_api.learner_exclusions import is_excluded_learner
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET

from .common import CONN, db_is_read_only
from .contract_documents import ensure_contract_archive_table, ensure_contract_uploads_table
from .evidence_documents import ensure_evidence_override_table

try:
    from learner_api.evidence_storage import download_blob_bytes
except ImportError:  # pragma: no cover - optional storage helper in slim test envs.
    download_blob_bytes = None


# --- small parsing helpers -------------------------------------------------

def _iso_date(value):
    """Pull a ``YYYY-MM-DD`` string out of the many timestamp formats used in
    the JSON. Returns None when nothing date-like is present."""
    if not value:
        return None
    match = re.search(r"(\d{4})-(\d{2})-(\d{2})", str(value))
    return match.group(0) if match else None


def _training_plan_from_audit(raw_plan):
    """Normalise the deployed ``Audit.learner_match.aptem_training_plan``."""
    import copy as _copy

    if isinstance(raw_plan, str):
        try:
            raw_plan = json.loads(raw_plan)
        except ValueError:
            raw_plan = []
    if not isinstance(raw_plan, list):
        raw_plan = []

    source_plan = _copy.deepcopy(raw_plan)
    months = []
    total = 0
    completed = 0
    for month in raw_plan:
        if not isinstance(month, dict):
            continue
        modules = []
        for item in month.get("modules") or []:
            if not isinstance(item, dict):
                continue
            component = item.get("components") if isinstance(item.get("components"), dict) else {}
            status = component.get("status") or "Unknown"
            total += 1
            if str(status).strip().lower() == "completed":
                completed += 1
            modules.append({
                "name": item.get("module") or "Untitled module",
                "type": component.get("type") or "",
                "status": status,
                "components": _copy.deepcopy(component),
                "raw": _copy.deepcopy(item),
            })
        months.append({
            "month": month.get("month") or "",
            "date": _iso_date(month.get("date")),
            "modules": modules,
            "raw": _copy.deepcopy(month),
        })

    return {
        "total_modules": total,
        "completed_modules": completed,
        "months": months,
        "raw": source_plan,
    }


def _learner_id(name):
    """Stable slug used as the learner filter key (id == lowercased name)."""
    return (name or "").strip().lower()


# --- contract signature-date extraction ------------------------------------

def _contract_blob_from_azure_path(value):
    """Parse `az://account/container/blob` contract references."""
    text = str(value or "").strip()
    if not text:
        return None
    parsed = urlparse(text)
    if parsed.scheme == "az":
        account = (parsed.netloc or "").strip()
        path = parsed.path.lstrip("/")
        if "/" not in path:
            return None
        container, blob_name = path.split("/", 1)
        configured_account = getattr(settings, "AZURE_STORAGE_ACCOUNT", "")
        if configured_account and account and account.lower() != configured_account.lower():
            return None
        return container, unquote(blob_name)
    if parsed.scheme == "https" and parsed.hostname:
        configured_account = getattr(settings, "AZURE_STORAGE_ACCOUNT", "")
        expected_host = f"{configured_account}.blob.core.windows.net".lower()
        if configured_account and parsed.hostname.lower() != expected_host:
            return None
        path = parsed.path.lstrip("/")
        if "/" not in path:
            return None
        container, blob_name = path.split("/", 1)
        return container, unquote(blob_name)
    return None


def _safe_contract_filename(document_name, blob_name):
    blob_filename = unquote(blob_name.rsplit("/", 1)[-1]) if blob_name else ""
    filename = blob_filename or str(document_name or "contract").strip() or "contract"
    filename = re.sub(r'[\r\n"\\]+', " ", filename).strip()
    return filename or "contract"


_CONTRACT_SIGNATURE_CACHE_TTL_SECONDS = 600
_contract_signature_cache = {}


def _normalise_contract_date(value):
    text = str(value or "").strip()
    if not text:
        return None
    iso_match = re.search(r"(?<!\d)(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})(?!\d)", text)
    if iso_match:
        year, month, day = (int(part) for part in iso_match.groups())
    else:
        uk_match = re.search(r"(?<!\d)(\d{1,2})[-/.](\d{1,2})[-/.](\d{2,4})(?!\d)", text)
        if not uk_match:
            return None
        day, month, year = (int(part) for part in uk_match.groups())
        if year < 100:
            year += 2000
    try:
        return datetime.date(year, month, day).isoformat()
    except ValueError:
        return None


def _extract_contract_text(filename, data):
    extension = "." + str(filename or "").rsplit(".", 1)[-1].lower() if "." in str(filename or "") else ""
    if extension == ".pdf":
        try:
            import fitz

            document = fitz.open(stream=data, filetype="pdf")
            lines = []
            for page_index in range(min(len(document), 30)):
                page = document[page_index]
                text = re.sub(r"\s+", " ", page.get_text("text") or "").strip()
                if text:
                    lines.append(f"Page {page_index + 1}: {text}")
            text = "\n".join(lines)
            if text:
                return text
        except Exception:
            pass
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            lines = []
            for page in reader.pages[:10]:
                text = re.sub(r"\s+", " ", page.extract_text() or "").strip()
                if text:
                    lines.append(text)
            return "\n".join(lines)
        except Exception:
            return ""
    if extension == ".docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            lines = []
            for paragraph in document.paragraphs:
                text = re.sub(r"\s+", " ", paragraph.text).strip()
                if text:
                    lines.append(text)
            for table in document.tables:
                for row in table.rows:
                    values = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        lines.append(" | ".join(values))
            return "\n".join(lines)
        except Exception:
            return ""
    if extension in {".txt", ".csv"}:
        return data.decode("utf-8-sig", errors="replace")
    return ""


def _contract_signature_dates_from_text(text):
    clean = re.sub(r"\s+", " ", str(text or "")).strip()
    if not clean:
        return {"learner_signed_date": None, "fully_signed_date": None}

    date_pattern = r"(\d{4}[-/.]\d{1,2}[-/.]\d{1,2}|\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})"
    signature_section = clean[-3500:]
    markers = []
    for marker_pattern in (
        r"\bsignatures?\s*&\s*declarations?\b",
        r"\bsignatories\b",
        r"\bsignatures\b",
    ):
        markers.extend(re.finditer(marker_pattern, clean, flags=re.IGNORECASE))
    if markers:
        marker = max(markers, key=lambda item: item.start())
        signature_section = clean[marker.start(): marker.start() + 3500]

    learner_date = None
    for pattern in (
        rf"(?:apprentice|learner)\s*:?.{{0,220}}?\bdate\s*:?\s*{date_pattern}",
        rf"(?:apprentice|learner)\s*:?.{{0,220}}?{date_pattern}",
    ):
        match = re.search(pattern, signature_section, flags=re.IGNORECASE)
        if match:
            learner_date = _normalise_contract_date(match.group(1))
            if learner_date:
                break

    all_dates = [
        parsed for parsed in (_normalise_contract_date(match.group(1)) for match in re.finditer(date_pattern, signature_section))
        if parsed
    ]
    return {
        "learner_signed_date": learner_date,
        "fully_signed_date": max(all_dates) if all_dates else None,
    }


_NO_SIGNATURE_DATES = {"learner_signed_date": None, "fully_signed_date": None}


def _ensure_signature_cache_table(cur):
    if db_is_read_only(cur):
        return
    cur.execute(
        '''
        create table if not exists "Manual_audit".contract_signature_cache (
            azure_path text primary key,
            learner_signed_date text,
            fully_signed_date text,
            extracted_at timestamp with time zone not null default now()
        )
        '''
    )


def _download_signature_dates(azure_path, document_name):
    """Download one contract blob and extract its signature dates.

    Returns (result, persistable): persistable is False on a download/parse
    failure so a transient Azure error never lands in the durable cache.
    """
    location = _contract_blob_from_azure_path(azure_path)
    if not location or download_blob_bytes is None:
        return dict(_NO_SIGNATURE_DATES), False
    try:
        container, blob_name = location
        filename = _safe_contract_filename(document_name, blob_name)
        data = download_blob_bytes(container, blob_name)
        text = _extract_contract_text(filename, data)
        return _contract_signature_dates_from_text(text), True
    except Exception:
        return dict(_NO_SIGNATURE_DATES), False


def _contract_signature_dates_many(cursor, items):
    """Resolve signature dates for many (azure_path, document_name) pairs.

    Layered lookup: in-process cache -> "Manual_audit".contract_signature_cache
    (a signed document at a fixed azure_path is immutable, so durable hits never
    expire) -> parallel Azure downloads for the misses only. The profile page
    used to take ~15s per cold load because every contract PDF was downloaded
    sequentially with only a 10-minute in-process cache.
    """
    wanted = {}
    for azure_path, document_name in items:
        key = str(azure_path or "")
        if key and key not in wanted and _contract_blob_from_azure_path(azure_path):
            wanted[key] = (azure_path, document_name)
    results = {}
    if not wanted:
        return results

    now = time.monotonic()
    for key in list(wanted):
        cached = _contract_signature_cache.get(key)
        if cached and now < cached["expires_at"]:
            results[key] = cached["value"]
            del wanted[key]
    if not wanted:
        return results

    _ensure_signature_cache_table(cursor)
    cursor.execute(
        '''
        select azure_path, learner_signed_date, fully_signed_date
        from "Manual_audit".contract_signature_cache
        where azure_path = any(%s)
        ''',
        [list(wanted)],
    )
    for azure_path, learner_signed, fully_signed in cursor.fetchall():
        value = {"learner_signed_date": learner_signed, "fully_signed_date": fully_signed}
        results[azure_path] = value
        _contract_signature_cache[azure_path] = {
            "expires_at": now + _CONTRACT_SIGNATURE_CACHE_TTL_SECONDS,
            "value": value,
        }
        wanted.pop(azure_path, None)
    if not wanted:
        return results

    if download_blob_bytes is None:
        for key in wanted:
            results[key] = dict(_NO_SIGNATURE_DATES)
        return results

    with ThreadPoolExecutor(max_workers=min(6, len(wanted))) as executor:
        downloads = {
            key: executor.submit(_download_signature_dates, azure_path, document_name)
            for key, (azure_path, document_name) in wanted.items()
        }
    for key, future in downloads.items():
        value, persistable = future.result()
        results[key] = value
        _contract_signature_cache[key] = {
            "expires_at": now + _CONTRACT_SIGNATURE_CACHE_TTL_SECONDS,
            "value": value,
        }
        if persistable:
            # Cache misses must never break the profile read: while Neon has
            # the endpoint read-only this insert fails, but the value is
            # already served from the in-process cache above.
            try:
                cursor.execute(
                    '''
                    insert into "Manual_audit".contract_signature_cache
                        (azure_path, learner_signed_date, fully_signed_date, extracted_at)
                    values (%s, %s, %s, now())
                    on conflict (azure_path) do update set
                        learner_signed_date = excluded.learner_signed_date,
                        fully_signed_date = excluded.fully_signed_date,
                        extracted_at = now()
                    ''',
                    [key, value["learner_signed_date"], value["fully_signed_date"]],
                )
            except DatabaseError:
                pass
    return results


# --- views -----------------------------------------------------------------

@require_GET
def health(_request: HttpRequest) -> JsonResponse:
    alias = CONN if CONN in connections.databases else "default"
    with connections[alias].cursor() as cursor:
        cursor.execute("SELECT current_database(), now()")
        database, timestamp = cursor.fetchone()
    return JsonResponse({"ok": True, "source": "Manual_audit", "database": database, "time": timestamp})


@require_GET
def learner_profile(request: HttpRequest) -> JsonResponse:
    """Return the rich cross-source profile for any Manual_audit learner."""
    learner_key = request.GET.get("learner", "").strip().lower()
    if not learner_key or len(learner_key) > 200:
        return JsonResponse({"error": "A valid learner is required."}, status=400)
    if is_excluded_learner(learner_key, learner_key):
        return JsonResponse({"error": "Learner not found."}, status=404)

    try:
        learner = _load_profile_learner(learner_key)
    except (KeyError, DatabaseError) as error:
        return JsonResponse(
            {"error": "Could not read the learner from Manual_audit.", "details": str(error)},
            status=503,
        )
    if learner is None:
        return JsonResponse({"error": "Learner not found."}, status=404)

    try:
        sources = _load_profile_sources(
            learner["aptem_id"],
            learner.get("email"),
            learner.get("name"),
        )
    except DatabaseError as error:
        return JsonResponse(
            {"error": "Could not load the learner profile sources.", "details": str(error)},
            status=503,
        )

    # Auditor-corrected dates replace the source values everywhere the profile
    # is consumed (journal header, profile page, PDF).
    try:
        with connections[CONN].cursor() as cur:
            _apply_profile_date_overrides(cur, learner["aptem_id"], sources["learning_delivery"])
    except (KeyError, DatabaseError):
        pass  # missing override table must never block the profile itself

    training_plan = _training_plan_from_audit(learner.get("training_plan"))

    return JsonResponse({
        "id": _learner_id(learner["name"]),
        "aptem_id": str(learner["aptem_id"]),
        "name": learner["name"],
        "email": learner.get("email"),
        "programme": learner.get("programme_name") or "Unknown programme",
        "programme_status": (
            sources["programme_status"]
            if sources["programme_status"] not in (None, "", "Unknown")
            else learner.get("programme_status") or "Unknown"
        ),
        "break_in_learning": sources["break_in_learning"],
        "coach": learner.get("coach") or {"name": None, "email": None},
        "planned_hours": sources["learning_delivery"].get("planned_hours"),
        "learning_delivery": sources["learning_delivery"],
        "contracts": sources["contracts"],
        "training_plan": training_plan,
        "skills_radar": sources["skills_radar"],
        "certifications": sources["certifications"],
        "employment": sources["employment"],
        "programme_understanding": sources["programme_understanding"],
    })


def _load_profile_learner(learner_key):
    """Resolve one learner from the Manual_audit mirror.

    ``Manual_audit.learners`` is this workspace's canonical learner list (a
    synced copy of ``Last_audit.learners``). The Training Plan is joined from
    the shared ``Audit.learner_match.aptem_training_plan`` by Aptem ID.
    """
    with connections[CONN].cursor() as cursor:
        cursor.execute(
            '''
            select l.aptem_id, l.learner_name, l.learner_email,
                   l.programme_name, l.programme_status,
                   l.coach_name, l.coach_email,
                   lm.aptem_training_plan
            from "Manual_audit".learners l
            left join "Audit".learner_match lm
              on lm.aptem_id = l.aptem_id
            where l.aptem_id::text = %s
               or lower(l.learner_name) = %s
            order by case when l.aptem_id::text = %s then 0 else 1 end
            limit 1
            ''',
            [learner_key, learner_key, learner_key],
        )
        row = cursor.fetchone()

    if row is None:
        return None

    (
        aptem_id, learner_name, learner_email, programme_name,
        programme_status, coach_name, coach_email, training_plan,
    ) = row
    if is_excluded_learner(aptem_id, learner_name):
        return None
    return {
        "aptem_id": aptem_id,
        "name": learner_name,
        "email": learner_email,
        "programme_name": programme_name,
        "programme_status": programme_status,
        "coach": {"name": coach_name, "email": coach_email},
        "training_plan": training_plan,
    }


def _contract_sort_key(entry):
    value = entry.get("date")
    if isinstance(value, (datetime.date, datetime.datetime)):
        return value.isoformat()
    return str(value or "")


def _load_profile_sources(aptem_id, learner_email, learner_name=None):
    """Read the profile's external sources without retaining user selection.

    Reference reads (contracts probe, skills radar, certifications, ILR,
    evidence) come from the shared source tables; every auditor write is read
    back from the ``Manual_audit`` tables this app owns.
    """
    contracts = []
    skill_groups = {}
    certifications = []
    employment = None
    learning_delivery = {}
    programme_understanding = {
        "understanding_programme": None,
        "career_development_progression": None,
    }
    programme_status = "Unknown"
    break_in_learning = {
        "has_break_in_learning": False,
        "last_learning_date": None,
        "expected_return_date": None,
        "has_return_to_learning": False,
        "return_to_learning_date": None,
        "revised_learning_planned_end_date": None,
    }

    with connections[CONN].cursor() as cursor:
        ensure_contract_archive_table(cursor)
        ensure_contract_uploads_table(cursor)
        cursor.execute(
            '''
            select contracts.id, coalesce(nullif(archive.display_name, ''), contracts.document_name), contracts.status, contracts.date,
                   contracts.learner_signed_date, contracts.fully_signed_date,
                   contracts.requested_date, contracts.program_name,
                   contracts.program_start_date, contracts.planned_end_date,
                   contracts.file, contracts.azure_path,
                   archive.archived_at, archive.archived_by
            from fetching_evidence.aptem_cv_contracts_probe contracts
            left join "Manual_audit".contract_document_archive archive
              on archive.contract_id = contracts.id
            where contracts.learner_id = %s
              and archive.deleted_at is null
            order by contracts.date desc nulls last, contracts.id desc
            ''',
            [aptem_id],
        )
        probe_rows = cursor.fetchall()

        # Manual uploads live in this workspace's own table, so they never leak
        # into the automatic audit workspace.
        cursor.execute(
            '''
            select id, coalesce(nullif(display_name, ''), document_name), status, date,
                   programme, azure_path, archived_at, archived_by
            from "Manual_audit".contract_uploads
            where learner_id = %s and deleted_at is null
            order by date desc nulls last, id desc
            ''',
            [aptem_id],
        )
        upload_rows = cursor.fetchall()

        # Resolve every contract's document-extracted signature dates in one
        # batched, cached, parallel pass (per-row sequential Azure downloads
        # made the profile take ~15s on a cold load).
        signature_map = _contract_signature_dates_many(
            cursor,
            [(row[11], row[1]) for row in probe_rows] + [(row[5], row[1]) for row in upload_rows],
        )

        for row in probe_rows:
            signature_dates = signature_map.get(str(row[11] or ""), _NO_SIGNATURE_DATES)
            document_learner_signed_date = signature_dates.get("learner_signed_date")
            document_fully_signed_date = signature_dates.get("fully_signed_date")
            learner_signed_date = document_learner_signed_date or row[4]
            fully_signed_date = document_fully_signed_date or row[5]
            contracts.append({
                "id": str(row[0]),
                "document_name": row[1] or "Contract",
                "status": row[2] or "Unknown",
                "date": row[3],
                "learner_signed_date": learner_signed_date,
                "fully_signed_date": fully_signed_date,
                "document_learner_signed_date": document_learner_signed_date,
                "document_fully_signed_date": document_fully_signed_date,
                "metadata_learner_signed_date": row[4],
                "metadata_fully_signed_date": row[5],
                "learner_signed_date_source": "document" if document_learner_signed_date else "metadata",
                "fully_signed_date_source": "document" if document_fully_signed_date else "metadata",
                "requested_date": row[6],
                "programme": row[7],
                "programme_start_date": row[8],
                "planned_end_date": row[9],
                "file": f"/manual_audit_api/contracts/{row[0]}/open" if row[11] else row[10],
                "download_file": row[10],
                "azure_path_available": bool(row[11]),
                "archived": bool(row[12]),
                "archived_at": row[12],
                "archived_by": row[13],
            })

        for row in upload_rows:
            signature_dates = signature_map.get(str(row[5] or ""), _NO_SIGNATURE_DATES)
            contracts.append({
                "id": f"manual-{row[0]}",
                "document_name": row[1] or "Contract",
                "status": row[2] or "Uploaded",
                "date": row[3],
                "learner_signed_date": signature_dates.get("learner_signed_date"),
                "fully_signed_date": signature_dates.get("fully_signed_date"),
                "document_learner_signed_date": signature_dates.get("learner_signed_date"),
                "document_fully_signed_date": signature_dates.get("fully_signed_date"),
                "metadata_learner_signed_date": None,
                "metadata_fully_signed_date": None,
                "learner_signed_date_source": "document" if signature_dates.get("learner_signed_date") else "metadata",
                "fully_signed_date_source": "document" if signature_dates.get("fully_signed_date") else "metadata",
                "requested_date": None,
                "programme": row[4],
                "programme_start_date": None,
                "planned_end_date": None,
                "file": f"/manual_audit_api/contracts/manual-{row[0]}/open" if row[5] else None,
                "download_file": None,
                "azure_path_available": bool(row[5]),
                "archived": bool(row[6]),
                "archived_at": row[6],
                "archived_by": row[7],
            })
        contracts.sort(key=_contract_sort_key, reverse=True)

        cursor.execute(
            '''
            select program_status, "Break in learning"
            from fetching_evidence.aptem_cv_contracts_probe
            where learner_id = %s and source <> 'audit_upload'
            order by fetched_at desc nulls last, id desc
            limit 1
            ''',
            [aptem_id],
        )
        status_row = cursor.fetchone()
        if status_row:
            programme_status = status_row[0] or "Unknown"
            break_value = status_row[1]
            if isinstance(break_value, str):
                try:
                    break_value = json.loads(break_value)
                except ValueError:
                    break_value = None
            if isinstance(break_value, dict):
                break_in_learning = {
                    "has_break_in_learning": bool(break_value.get("has_break_in_learning")),
                    "last_learning_date": break_value.get("last_learning_date"),
                    "expected_return_date": break_value.get("expected_return_date"),
                    "has_return_to_learning": bool(break_value.get("has_return_to_learning")),
                    "return_to_learning_date": break_value.get("return_to_learning_date"),
                    "revised_learning_planned_end_date": break_value.get("revised_learning_planned_end_date"),
                }
            if str(programme_status).strip().lower() == "onbreak":
                break_in_learning["has_break_in_learning"] = True

        cursor.execute(
            '''
            select
                coalesce(
                    "Programme understanding" ->> 'understanding_programme',
                    "Programme understanding" -> 'raw' ->> 'ExtendedILRModel_UnderstandingProgramme'
                ),
                coalesce(
                    "Programme understanding" ->> 'career_development_progression',
                    "Programme understanding" -> 'raw' ->> 'ExtendedILRModel_CareerDevelopmentProgression'
                )
            from fetching_evidence.aptem_cv_contracts_probe
            where learner_id = %s
              and "Programme understanding" is not null
            order by fetched_at desc nulls last, id desc
            limit 1
            ''',
            [aptem_id],
        )
        understanding_row = cursor.fetchone()
        if understanding_row:
            programme_understanding = {
                "understanding_programme": understanding_row[0] or None,
                "career_development_progression": understanding_row[1] or None,
            }

        cursor.execute(
            '''
            select characteristic_name, assessed_level
            from fetching_evidence.aptem_skills_radar_probe
            where learner_id = %s and assessed_level is not null
            order by characteristic_name
            ''',
            [aptem_id],
        )
        for characteristic, assessed_level in cursor.fetchall():
            name = (characteristic or "").strip()
            score = max(0, min(8, int(assessed_level)))
            # Two source formats coexist in aptem_skills_radar_probe:
            #  - Project Controls: "Understanding of X (Knowledge) - K7: ..."
            #  - Marketing:        "K3: I understand ..." / "S1: I can ..." / "B6: Acting ..."
            # Classifying by the (Knowledge|Skill|Behaviour) suffix alone dumped
            # every Marketing row into Skills, leaving the profile's Knowledge
            # and Behaviours groups empty for those programmes.
            legacy = re.match(r"Understanding of (.+?) \((Knowledge|Skill|Behaviour)\)", name)
            coded = re.match(r"([KSB])\d+\s*[:.\-]", name, flags=re.IGNORECASE)
            if legacy:
                domain = legacy.group(1).strip()
                score_type = legacy.group(2).lower()
            elif coded:
                domain = name.split(" - ")[0].strip()
                score_type = {"K": "knowledge", "S": "skill", "B": "behaviour"}[coded.group(1).upper()]
            else:
                domain = name.split(" - ")[0].strip() or "Skill"
                score_type = "skill"
            field = {"knowledge": "knowledge", "skill": "skill_score", "behaviour": "behaviour"}[score_type]
            skill_groups.setdefault(domain, {})[field] = score

        cursor.execute(
            '''
            select certifications, employment_details
            from fetching_evidence.aptem_cv_certifications
            where learner_id = %s
            order by updated_at desc nulls last, id desc
            ''',
            [aptem_id],
        )
        seen_certifications = set()
        for certification_value, employment_value in cursor.fetchall():
            if isinstance(certification_value, str):
                try:
                    certification_value = json.loads(certification_value)
                except ValueError:
                    certification_value = []
            if isinstance(employment_value, str):
                try:
                    employment_value = json.loads(employment_value)
                except ValueError:
                    employment_value = []
            if isinstance(certification_value, list):
                for certification in certification_value:
                    if not isinstance(certification, dict):
                        continue
                    key = (
                        str(certification.get("name") or "").strip().lower(),
                        str(certification.get("issuer") or "").strip().lower(),
                    )
                    if key in seen_certifications or not key[0]:
                        continue
                    seen_certifications.add(key)
                    certifications.append(certification)
            if employment is None:
                employment = _first_employment_details(employment_value)

        if learner_email or learner_name:
            learner_email = str(learner_email or "").strip()
            learner_name = str(learner_name or "").strip()
            cursor.execute(
                '''
                select learn_ref_number, planned_hours, otj_actual_hours,
                       learn_start_date, learn_plan_end_date, completion_status,
                       nullif(
                           concat_ws(', ',
                               nullif(btrim(address_line_1), ''),
                               nullif(btrim(address_line_2), ''),
                               nullif(btrim(address_line_3), ''),
                               nullif(btrim(address_line_4), '')
                           ),
                           ''
                       ) as learner_address,
                       nullif(btrim(postcode), '') as learner_postcode,
                       nullif(btrim(delivery_location_postcode), '') as employer_postcode
                from "Audit".ilr_learning_deliveries
                where planned_hours is not null
                  and (
                    (%s <> '' and lower(btrim(email)) = lower(%s))
                    or
                    (%s <> '' and lower(btrim(concat_ws(' ', given_names, family_name))) = lower(%s))
                  )
                order by
                    case when %s <> '' and lower(btrim(email)) = lower(%s) then 0 else 1 end,
                    aim_seq_number, updated_at desc nulls last, id desc
                limit 1
                ''',
                [
                    learner_email, learner_email,
                    learner_name, learner_name,
                    learner_email, learner_email,
                ],
            )
            delivery = cursor.fetchone()
            if delivery:
                learning_delivery = {
                    "learner_reference": delivery[0],
                    "planned_hours": delivery[1],
                    "actual_hours": delivery[2],
                    "start_date": delivery[3],
                    "planned_end_date": delivery[4],
                    "completion_status": delivery[5],
                    "learner_address": delivery[6],
                    "learner_postcode": delivery[7],
                    "employer_postcode": delivery[8],
                    "first_evidence_date": None,
                    "first_evidence_items": [],
                    "archived_evidence_items": [],
                }
                ensure_evidence_override_table(cursor)
                cursor.execute(
                    '''
                    with raw_candidates as (
                        select
                            item ->> 'id' as evidence_id,
                            item ->> 'name' as evidence_name,
                            item ->> 'component_name' as component_name,
                            item ->> 'kind' as evidence_kind,
                            item ->> 'status' as evidence_status,
                            item ->> 'file' as evidence_file,
                            item ->> 'content' as evidence_content,
                            substring(item ->> 'created_date' from 1 for 10)::date as evidence_date
                        from fetching_evidence.learner_evidence learner_evidence
                        cross join lateral jsonb_array_elements(
                            case
                                when jsonb_typeof(learner_evidence.evidence) = 'array'
                                    then learner_evidence.evidence
                                else '[]'::jsonb
                            end
                        ) item
                        where learner_evidence.learner_id = %s
                          and ltrim(lower(coalesce(item ->> 'name', ''))) not like 'welcome%%'
                          and ltrim(lower(coalesce(item ->> 'component_name', ''))) not like 'welcome%%'
                          and coalesce(item ->> 'created_date', '') ~ '^\\d{4}-\\d{2}-\\d{2}'
                          and substring(item ->> 'created_date' from 1 for 10)::date >= %s
                    ), candidates as (
                        select distinct on (evidence_id)
                            evidence_id, evidence_name, component_name, evidence_kind,
                            evidence_status, evidence_file, evidence_content, evidence_date
                        from raw_candidates
                        order by evidence_id, evidence_date
                    )
                    select candidates.evidence_id, candidates.evidence_name,
                           candidates.component_name, candidates.evidence_kind,
                           candidates.evidence_status, candidates.evidence_file,
                           candidates.evidence_content,
                           coalesce(overrides.evidence_date, candidates.evidence_date) as evidence_date,
                           overrides.archived_at is not null as archived,
                           overrides.deleted_at is not null as deleted,
                           false as uploaded
                    from candidates
                    left join "Manual_audit".learner_evidence_overrides overrides
                      on overrides.learner_id = %s
                     and overrides.is_uploaded = false
                     and overrides.source_evidence_id::text = candidates.evidence_id
                    union all
                    select uploads.evidence_id, uploads.document_name,
                           uploads.component_name, uploads.evidence_kind,
                           uploads.evidence_status, null, null, uploads.evidence_date,
                           uploads.archived_at is not null as archived,
                           uploads.deleted_at is not null as deleted,
                           true as uploaded
                    from "Manual_audit".learner_evidence_overrides uploads
                    where uploads.learner_id = %s and uploads.is_uploaded = true
                    order by evidence_date, evidence_id
                    ''',
                    [aptem_id, delivery[3], aptem_id, aptem_id],
                )
                evidence_rows = cursor.fetchall()
                if evidence_rows:
                    evidence_items = [
                        {
                            "id": row[0],
                            "name": row[1] or "Untitled evidence",
                            "component_name": row[2] or "",
                            "kind": row[3] or "",
                            "status": row[4] or "",
                            "file": f"/manual_audit_api/evidence/{quote(str(row[0]), safe='')}/open?learner_id={aptem_id}" if row[0] else None,
                            "content": row[6],
                            "date": row[7],
                            "archived": bool(row[8]),
                            "deleted": bool(row[9]),
                            "uploaded": bool(row[10]),
                        }
                        for row in evidence_rows
                    ]
                    first_date, first_items, archived_items = _partition_evidence_items(evidence_items)
                    learning_delivery["archived_evidence_items"] = archived_items
                    learning_delivery["first_evidence_date"] = first_date
                    learning_delivery["first_evidence_items"] = first_items

    skills_radar = [
        {
            "skill": domain,
            "knowledge": scores.get("knowledge"),
            "skill_score": scores.get("skill_score"),
            "behaviour": scores.get("behaviour"),
            "maximum": 8,
        }
        for domain, scores in sorted(skill_groups.items())
    ]
    return {
        "contracts": contracts,
        "skills_radar": skills_radar,
        "certifications": certifications,
        "employment": employment,
        "learning_delivery": learning_delivery,
        "programme_understanding": programme_understanding,
        "programme_status": programme_status,
        "break_in_learning": break_in_learning,
    }


def _partition_evidence_items(evidence_items):
    """Keep Aptem's original first evidence fixed; uploads may replace it."""
    archived_items = [
        item for item in evidence_items if item["archived"] and not item["deleted"]
    ]
    source_items = [item for item in evidence_items if not item["uploaded"]]
    original_source_date = min((item["date"] for item in source_items), default=None)
    qualifying_items = [
        item for item in evidence_items
        if not item["archived"]
        and not item["deleted"]
        and (
            item["uploaded"]
            or (original_source_date is not None and item["date"] == original_source_date)
        )
    ]
    first_date = min((item["date"] for item in qualifying_items), default=None)
    first_items = [item for item in qualifying_items if item["date"] == first_date]
    return first_date, first_items, archived_items


def _first_employment_details(value):
    if isinstance(value, dict):
        nested = value.get("employment_details")
        if isinstance(nested, dict) and nested.get("section_found", True):
            return nested
        if value.get("employer_name") and value.get("section_found", True):
            return value
    if isinstance(value, list):
        for item in value:
            details = _first_employment_details(item)
            if details:
                return details
    return None


# --- manual auditor annotations (KSBs + planned hours) per activity ---------

def _ensure_annotation_table(cur):
    if db_is_read_only(cur):
        return
    cur.execute(
        '''
        create table if not exists "Manual_audit".activity_annotations (
            component_id text primary key,
            planned_hours numeric,
            mapped_ksbs text,
            updated_by text,
            updated_at timestamp with time zone default now()
        )
        '''
    )


def _annotation_payload(row):
    if not row:
        return {"planned_hours": None, "mapped_ksbs": None, "updated_by": None, "updated_at": None}
    component_id, planned_hours, mapped_ksbs, updated_by, updated_at = row
    return {
        "component_id": component_id,
        "planned_hours": float(planned_hours) if planned_hours is not None else None,
        "mapped_ksbs": mapped_ksbs,
        "updated_by": updated_by,
        "updated_at": updated_at.isoformat() if updated_at else None,
    }


@require_GET
def activity_annotation(request: HttpRequest) -> JsonResponse:
    """Return the auditor-entered KSBs / planned hours for one activity."""
    component = request.GET.get("component", "").strip()
    if not component:
        return JsonResponse({"error": "component is required"}, status=400)
    try:
        with connections[CONN].cursor() as cur:
            _ensure_annotation_table(cur)
            cur.execute(
                '''
                select component_id, planned_hours, mapped_ksbs, updated_by, updated_at
                from "Manual_audit".activity_annotations where component_id = %s
                ''',
                [component],
            )
            row = cur.fetchone()
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not read activity annotations.", "details": str(error)}, status=503)
    payload = _annotation_payload(row)
    payload.setdefault("component_id", component)
    return JsonResponse(payload)


@csrf_exempt
def save_activity_annotation(request: HttpRequest) -> JsonResponse:
    """Create/update the auditor-entered KSBs / planned hours for one activity."""
    if request.method != "POST":
        return JsonResponse({"error": "Method not allowed."}, status=405)
    try:
        body = json.loads(request.body or b"{}")
    except ValueError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    component = str(body.get("component_id") or "").strip()
    if not component:
        return JsonResponse({"error": "component_id is required"}, status=400)

    planned_raw = body.get("planned_hours")
    if planned_raw in (None, ""):
        planned_hours = None
    else:
        try:
            planned_hours = float(planned_raw)
        except (TypeError, ValueError):
            return JsonResponse({"error": "planned_hours must be a number"}, status=400)
        if not math.isfinite(planned_hours):
            return JsonResponse({"error": "planned_hours must be a finite number"}, status=400)
        if planned_hours < 0 or planned_hours > 100000:
            return JsonResponse({"error": "planned_hours is out of range"}, status=400)

    mapped_ksbs = body.get("mapped_ksbs")
    if mapped_ksbs is not None:
        mapped_ksbs = str(mapped_ksbs).strip()[:5000] or None
    updated_by = str(body.get("updated_by") or "").strip()[:200] or None

    try:
        with connections[CONN].cursor() as cur:
            _ensure_annotation_table(cur)
            cur.execute(
                '''
                insert into "Manual_audit".activity_annotations
                    (component_id, planned_hours, mapped_ksbs, updated_by, updated_at)
                values (%s, %s, %s, %s, now())
                on conflict (component_id) do update set
                    planned_hours = excluded.planned_hours,
                    mapped_ksbs = excluded.mapped_ksbs,
                    updated_by = excluded.updated_by,
                    updated_at = now()
                returning component_id, planned_hours, mapped_ksbs, updated_by, updated_at
                ''',
                [component, planned_hours, mapped_ksbs, updated_by],
            )
            row = cur.fetchone()
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not save activity annotation.", "details": str(error)}, status=503)
    return JsonResponse(_annotation_payload(row))


# --- manual per-learner planned/actual hours overrides -----------------------
#
# The Learner search table's Planned/Actual hours are aggregates computed from
# the synced mirror (ILR planned hours + mapped activity seconds). The auditor
# can replace either number per learner — for one month or for the all-months
# total (period '') — directly from that table. The replacements live in this
# manual-owned table so the mirror stays safe to truncate + re-sync.

_HOURS_PERIOD_RE = re.compile(r"^\d{4}-(0[1-9]|1[0-2])$")


def _ensure_learner_hours_table(cur):
    if db_is_read_only(cur):
        return
    cur.execute(
        '''
        create table if not exists "Manual_audit".learner_hours_overrides (
            aptem_id bigint not null,
            period text not null default '',
            planned_hours numeric,
            actual_hours numeric,
            not_accepted_hours numeric,
            updated_by text,
            updated_at timestamp with time zone not null default now(),
            primary key (aptem_id, period)
        )
        '''
    )
    # Existing deployments predate the not-accepted column.
    cur.execute(
        'alter table "Manual_audit".learner_hours_overrides '
        'add column if not exists not_accepted_hours numeric'
    )


def _hours_override_value(value, field):
    """None / '' clears the override; otherwise a finite number 0..100000."""
    if value in (None, ""):
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        raise ValueError(f"{field} must be a number")
    # NaN slips past plain range comparisons and would poison the JSON feed.
    if not math.isfinite(number):
        raise ValueError(f"{field} must be a finite number")
    if number < 0 or number > 100000:
        raise ValueError(f"{field} is out of range (0–100000 hours)")
    return round(number, 2)


def _hours_override_period(value):
    period = str(value or "").strip()
    if period and period != "undated" and not _HOURS_PERIOD_RE.match(period):
        raise ValueError("period must be '' (all months), 'undated', or YYYY-MM")
    return period


def _learner_hours_payload(row):
    aptem_id, period, planned, actual, not_accepted, updated_by, updated_at = row
    return {
        "aptem_id": aptem_id,
        "period": period,
        "planned_hours": float(planned) if planned is not None else None,
        "actual_hours": float(actual) if actual is not None else None,
        "not_accepted_hours": float(not_accepted) if not_accepted is not None else None,
        "updated_by": updated_by,
        "updated_at": updated_at.isoformat() if updated_at else None,
    }


@csrf_exempt
def learner_hours(request: HttpRequest) -> JsonResponse:
    """GET every override; POST upsert one learner+period; DELETE remove one.

    POST body: {aptem_id, period?, planned_hours?, actual_hours?, updated_by?}.
    A field that is absent keeps its stored value; null/'' clears it. When both
    hours end up cleared the override row is removed (falls back to the
    computed mirror values).
    """
    if request.method == "GET":
        try:
            with connections[CONN].cursor() as cur:
                _ensure_learner_hours_table(cur)
                cur.execute(
                    '''
                    select aptem_id, period, planned_hours, actual_hours,
                           not_accepted_hours, updated_by, updated_at
                    from "Manual_audit".learner_hours_overrides
                    order by aptem_id, period
                    '''
                )
                rows = cur.fetchall()
        except (KeyError, DatabaseError) as error:
            return JsonResponse({"error": "Could not read learner hours overrides.", "details": str(error)}, status=503)
        return JsonResponse({"items": [_learner_hours_payload(row) for row in rows]})

    if request.method not in {"POST", "DELETE"}:
        return JsonResponse({"error": "Method not allowed."}, status=405)

    try:
        body = json.loads(request.body or b"{}")
    except ValueError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)
    if not isinstance(body, dict):
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    raw_aptem_id = body.get("aptem_id")
    if isinstance(raw_aptem_id, bool):
        return JsonResponse({"error": "aptem_id is required"}, status=400)
    try:
        aptem_id = int(raw_aptem_id)
    except (TypeError, ValueError):
        return JsonResponse({"error": "aptem_id is required"}, status=400)
    if aptem_id <= 0:
        return JsonResponse({"error": "aptem_id is required"}, status=400)
    try:
        period = _hours_override_period(body.get("period"))
    except ValueError as error:
        return JsonResponse({"error": str(error)}, status=400)

    if request.method == "DELETE":
        try:
            with connections[CONN].cursor() as cur:
                _ensure_learner_hours_table(cur)
                cur.execute(
                    'delete from "Manual_audit".learner_hours_overrides where aptem_id = %s and period = %s',
                    [aptem_id, period],
                )
        except (KeyError, DatabaseError) as error:
            return JsonResponse({"error": "Could not delete the learner hours override.", "details": str(error)}, status=503)
        return JsonResponse({"ok": True, "aptem_id": aptem_id, "period": period, "deleted": True})

    updated_by = str(body.get("updated_by") or "").strip()[:200] or None
    set_planned = "planned_hours" in body
    set_actual = "actual_hours" in body
    set_not_accepted = "not_accepted_hours" in body
    try:
        planned = _hours_override_value(body.get("planned_hours"), "planned_hours") if set_planned else None
        actual = _hours_override_value(body.get("actual_hours"), "actual_hours") if set_actual else None
        not_accepted = (
            _hours_override_value(body.get("not_accepted_hours"), "not_accepted_hours")
            if set_not_accepted else None
        )
    except ValueError as error:
        return JsonResponse({"error": str(error)}, status=400)

    try:
        if _overlay_learner(aptem_id) is None:
            return JsonResponse({"error": "Learner is outside the manual-audit cohort."}, status=404)
        with transaction.atomic(using=CONN):
            with connections[CONN].cursor() as cur:
                _ensure_learner_hours_table(cur)
                # Single-statement merge: a field absent from the body keeps its
                # stored value, so two auditors editing different fields never
                # overwrite each other (no read-modify-write window).
                cur.execute(
                    '''
                    insert into "Manual_audit".learner_hours_overrides as overrides
                        (aptem_id, period, planned_hours, actual_hours, not_accepted_hours, updated_by, updated_at)
                    values (%s, %s, %s, %s, %s, %s, now())
                    on conflict (aptem_id, period) do update set
                        planned_hours = case when %s then excluded.planned_hours else overrides.planned_hours end,
                        actual_hours = case when %s then excluded.actual_hours else overrides.actual_hours end,
                        not_accepted_hours = case when %s then excluded.not_accepted_hours else overrides.not_accepted_hours end,
                        updated_by = excluded.updated_by,
                        updated_at = now()
                    returning aptem_id, period, planned_hours, actual_hours, not_accepted_hours, updated_by, updated_at
                    ''',
                    [aptem_id, period, planned, actual, not_accepted, updated_by,
                     set_planned, set_actual, set_not_accepted],
                )
                row = cur.fetchone()
                if row[2] is None and row[3] is None and row[4] is None:
                    # Everything cleared -> no override left, drop the row.
                    cur.execute(
                        'delete from "Manual_audit".learner_hours_overrides where aptem_id = %s and period = %s',
                        [aptem_id, period],
                    )
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not save the learner hours override.", "details": str(error)}, status=503)
    return JsonResponse(_learner_hours_payload(row))


# --- manual per-learner profile-date overrides --------------------------------
#
# Start / first-evidence / planned-end dates come from read-only shared sources
# (ILR + fetched evidence). The auditor can replace any of them per learner;
# the replacement is applied whenever the profile is served, so the journal
# header, the profile page, and the PDF all show the corrected date.

_PROFILE_DATE_FIELDS = ("start_date", "first_evidence_date", "planned_end_date")


def _ensure_profile_dates_table(cur):
    if db_is_read_only(cur):
        return
    cur.execute(
        '''
        create table if not exists "Manual_audit".learner_profile_date_overrides (
            aptem_id bigint primary key,
            start_date date,
            first_evidence_date date,
            planned_end_date date,
            updated_by text,
            updated_at timestamp with time zone not null default now()
        )
        '''
    )


def _profile_date_value(value, field):
    """None / '' clears the override; otherwise a strict ISO date."""
    if value in (None, ""):
        return None
    try:
        return datetime.date.fromisoformat(str(value).strip())
    except ValueError:
        raise ValueError(f"{field} must be YYYY-MM-DD")


def _profile_dates_payload(row):
    aptem_id, start, first_evidence, planned_end, updated_by, updated_at = row
    return {
        "aptem_id": aptem_id,
        "start_date": start.isoformat() if start else None,
        "first_evidence_date": first_evidence.isoformat() if first_evidence else None,
        "planned_end_date": planned_end.isoformat() if planned_end else None,
        "updated_by": updated_by,
        "updated_at": updated_at.isoformat() if updated_at else None,
    }


def _apply_profile_date_overrides(cur, aptem_id, learning_delivery):
    """Overlay any auditor-corrected dates onto the learning_delivery dict."""
    _ensure_profile_dates_table(cur)
    cur.execute(
        '''
        select start_date, first_evidence_date, planned_end_date
        from "Manual_audit".learner_profile_date_overrides where aptem_id = %s
        ''',
        [aptem_id],
    )
    row = cur.fetchone()
    if not row:
        return
    for field, value in zip(_PROFILE_DATE_FIELDS, row):
        if value is not None:
            learning_delivery[field] = value.isoformat()
            learning_delivery[f"{field}_overridden"] = True


@csrf_exempt
def learner_profile_dates(request: HttpRequest) -> JsonResponse:
    """POST {aptem_id, start_date?, first_evidence_date?, planned_end_date?}.

    Absent field keeps its stored value; null/'' clears it (falls back to the
    source date). The row is dropped when everything is cleared.
    """
    if request.method != "POST":
        return JsonResponse({"error": "Method not allowed."}, status=405)
    try:
        body = json.loads(request.body or b"{}")
    except ValueError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)
    if not isinstance(body, dict):
        return JsonResponse({"error": "Invalid JSON body."}, status=400)
    raw_aptem_id = body.get("aptem_id")
    if isinstance(raw_aptem_id, bool):
        return JsonResponse({"error": "aptem_id is required"}, status=400)
    try:
        aptem_id = int(raw_aptem_id)
    except (TypeError, ValueError):
        return JsonResponse({"error": "aptem_id is required"}, status=400)
    if aptem_id <= 0:
        return JsonResponse({"error": "aptem_id is required"}, status=400)

    updated_by = str(body.get("updated_by") or "").strip()[:200] or None
    flags, values = [], []
    try:
        for field in _PROFILE_DATE_FIELDS:
            flags.append(field in body)
            values.append(_profile_date_value(body.get(field), field) if field in body else None)
    except ValueError as error:
        return JsonResponse({"error": str(error)}, status=400)

    try:
        if _overlay_learner(aptem_id) is None:
            return JsonResponse({"error": "Learner is outside the manual-audit cohort."}, status=404)
        with transaction.atomic(using=CONN):
            with connections[CONN].cursor() as cur:
                _ensure_profile_dates_table(cur)
                cur.execute(
                    '''
                    insert into "Manual_audit".learner_profile_date_overrides as overrides
                        (aptem_id, start_date, first_evidence_date, planned_end_date, updated_by, updated_at)
                    values (%s, %s, %s, %s, %s, now())
                    on conflict (aptem_id) do update set
                        start_date = case when %s then excluded.start_date else overrides.start_date end,
                        first_evidence_date = case when %s then excluded.first_evidence_date else overrides.first_evidence_date end,
                        planned_end_date = case when %s then excluded.planned_end_date else overrides.planned_end_date end,
                        updated_by = excluded.updated_by,
                        updated_at = now()
                    returning aptem_id, start_date, first_evidence_date, planned_end_date, updated_by, updated_at
                    ''',
                    [aptem_id, *values, updated_by, *flags],
                )
                row = cur.fetchone()
                if row[1] is None and row[2] is None and row[3] is None:
                    cur.execute(
                        'delete from "Manual_audit".learner_profile_date_overrides where aptem_id = %s',
                        [aptem_id],
                    )
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not save the profile dates.", "details": str(error)}, status=503)
    return JsonResponse(_profile_dates_payload(row))


# --- manual activity create/replace/delete overlay --------------------------
#
# The Manual_audit mirror is an immutable synced source. New auditor rows,
# edits and deletions are kept as an overlay so the mirror can be re-synced at
# any time without losing manual work. Deletes are soft deletes (tombstones).

_OVERLAY_CATEGORIES = {"attendance", "assignment", "video", "audio", "reading+quiz"}


def _ensure_activity_overlay_table(cur):
    if db_is_read_only(cur):
        return
    cur.execute(
        '''
        create table if not exists "Manual_audit".activity_overrides (
            aptem_id bigint not null,
            activity_id text not null,
            operation text not null check (operation in ('created', 'deleted', 'replaced')),
            payload jsonb not null,
            source_payload jsonb,
            updated_by text,
            created_at timestamp with time zone not null default now(),
            updated_at timestamp with time zone not null default now(),
            primary key (aptem_id, activity_id)
        )
        '''
    )


def _overlay_number(value, field):
    if value in (None, ""):
        return 0.0
    try:
        number = float(value)
    except (TypeError, ValueError):
        raise ValueError(f"{field} must be a number")
    # NaN slips past plain range comparisons and would poison the JSON feed.
    if not math.isfinite(number):
        raise ValueError(f"{field} must be a finite number")
    if number < 0 or number > 50:
        raise ValueError(f"{field} must be between 0 and 50 hours")
    return round(number, 4)


def _overlay_timestamp(value, field):
    if value in (None, ""):
        return None
    try:
        parsed = datetime.datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        raise ValueError(f"{field} must be an ISO timestamp")
    return parsed.isoformat()


def _validate_overlay_activity(raw, *, aptem_id, learner_name, activity_id, allow_any_category=False):
    if not isinstance(raw, dict):
        raise ValueError("activity must be an object")
    date = str(raw.get("date") or "").strip()
    try:
        parsed_date = datetime.date.fromisoformat(date)
    except ValueError:
        raise ValueError("date must use YYYY-MM-DD format")
    category = str(raw.get("category") or "").strip().lower()
    if category not in _OVERLAY_CATEGORIES:
        # Edits of MIRROR rows must not be blocked (or silently recategorised)
        # by source spellings outside the canonical five — e.g. "reading",
        # "quiz", or the "activity" fallback. New audit-created rows stay
        # strict: their category comes from the UI select.
        if not (allow_any_category and re.fullmatch(r"[a-z0-9+][a-z0-9+ _-]{0,49}", category)):
            raise ValueError(f"category must be one of: {', '.join(sorted(_OVERLAY_CATEGORIES))}")
    name = str(raw.get("activity") or "").strip()
    if not name:
        raise ValueError("activity is required")
    if len(name) > 500:
        raise ValueError("activity must be at most 500 characters")

    planned = _overlay_number(raw.get("planned"), "planned")
    actual = _overlay_number(raw.get("actual"), "actual")
    started = _overlay_timestamp(raw.get("timestamp_from"), "timestamp_from")
    completed = _overlay_timestamp(raw.get("timestamp_to"), "timestamp_to")
    if started and completed:
        start_value = datetime.datetime.fromisoformat(started)
        end_value = datetime.datetime.fromisoformat(completed)
        if start_value.tzinfo is None and end_value.tzinfo is not None or start_value.tzinfo is not None and end_value.tzinfo is None:
            raise ValueError("timestamps must use the same timezone style")
        if end_value < start_value:
            raise ValueError("timestamp_to must be after timestamp_from")

    display = str(raw.get("timestamp_display") or "").strip()[:100]
    if not display:
        if started and completed:
            display = f"{datetime.datetime.fromisoformat(started):%H:%M}–{datetime.datetime.fromisoformat(completed):%H:%M}"
        elif actual > 0:
            display = "input"

    return {
        "activity_id": activity_id,
        "learner_id": aptem_id,
        "learner_name": learner_name,
        "date": parsed_date.isoformat(),
        "month": parsed_date.strftime("%Y-%m"),
        "month_label": parsed_date.strftime("%B %Y"),
        "category": category,
        "activity": name,
        "activity_subtitle": str(raw.get("activity_subtitle") or "").strip()[:2000] or None,
        "planned": planned,
        "actual": actual,
        "timestamp_from": started,
        "timestamp_to": completed,
        "timestamp_display": display,
        "completed": bool(raw.get("completed", actual > 0)),
        "ksbs": raw.get("ksbs") if isinstance(raw.get("ksbs"), dict) else {"K": [], "S": [], "B": []},
        "iframe_url": None,
        "not_accepted": bool(raw.get("not_accepted", False)),
        "reporting_week_label": str(raw.get("reporting_week_label") or "").strip()[:200] or None,
        "audit_created": bool(raw.get("audit_created", str(activity_id).startswith("audit:"))),
        # Soft reference to an uploaded manual-audit evidence file (assignment
        # uploads); served by /manual_audit_api/evidence/<id>/open.
        "evidence_id": str(raw.get("evidence_id") or "").strip()[:100] or None,
        "evidence_name": str(raw.get("evidence_name") or "").strip()[:300] or None,
    }


def _overlay_learner(aptem_id):
    """Resolve a writable learner from the Manual_audit mirror by Aptem ID."""
    with connections[CONN].cursor() as cur:
        cur.execute(
            '''
            select learner_name
            from "Manual_audit".learners
            where aptem_id = %s
            limit 1
            ''',
            [aptem_id],
        )
        row = cur.fetchone()
    if not row:
        return None
    if is_excluded_learner(aptem_id, row[0]):
        return None
    return {"aptem_id": aptem_id, "name": row[0] or f"Learner {aptem_id}"}


@csrf_exempt
def activity_overrides(request: HttpRequest) -> JsonResponse:
    """List/create/update/soft-delete manual-audit activity overlays."""
    if request.method == "GET":
        raw_id = request.GET.get("aptem_id", "").strip()
        try:
            aptem_id = int(raw_id) if raw_id else None
        except ValueError:
            return JsonResponse({"error": "aptem_id must be an integer"}, status=400)
        try:
            with connections[CONN].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                if aptem_id is None:
                    cur.execute('''select aptem_id, activity_id, operation, payload, source_payload, updated_by, updated_at, created_at from "Manual_audit".activity_overrides order by updated_at''')
                else:
                    cur.execute('''select aptem_id, activity_id, operation, payload, source_payload, updated_by, updated_at, created_at from "Manual_audit".activity_overrides where aptem_id = %s order by updated_at''', [aptem_id])
                rows = cur.fetchall()
        except (KeyError, DatabaseError) as error:
            return JsonResponse({"error": "Could not read activity overrides.", "details": str(error)}, status=503)
        return JsonResponse({"items": [
            {"aptem_id": row[0], "activity_id": row[1], "operation": row[2], "payload": row[3], "source_payload": row[4], "updated_by": row[5], "updated_at": row[6].isoformat() if row[6] else None, "created_at": row[7].isoformat() if row[7] else None}
            for row in rows
        ]})

    if request.method not in {"POST", "PUT", "PATCH", "DELETE"}:
        return JsonResponse({"error": "Method not allowed."}, status=405)
    try:
        body = json.loads(request.body or b"{}")
        aptem_id = int(body.get("aptem_id"))
    except (TypeError, ValueError):
        return JsonResponse({"error": "aptem_id must be an integer"}, status=400)
    try:
        learner = _overlay_learner(aptem_id)
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not validate learner.", "details": str(error)}, status=503)
    if not learner:
        return JsonResponse({"error": "Learner is outside the manual-audit cohort."}, status=404)

    updated_by = str(body.get("updated_by") or "").strip()[:200] or None
    activity_id = str(body.get("activity_id") or "").strip()
    source_payload = None

    # POST with group_id: create the SAME activity once per member of that LMS
    # group (each learner tracks their own hours/completion, so per-learner
    # overlay rows are the shared-activity mechanism here).
    if request.method == "POST" and body.get("group_id") not in (None, ""):
        try:
            lms_group_id = int(body.get("group_id"))
        except (TypeError, ValueError):
            return JsonResponse({"error": "group_id must be an integer"}, status=400)
        try:
            with connections[CONN].cursor() as cur:
                cur.execute(
                    '''
                    select distinct l.aptem_id, l.learner_name
                    from "Manual_audit".group_learners gl
                    join "Manual_audit".learners l on l.learner_id = gl.learner_id
                    where gl.group_id = %s and l.aptem_id is not null
                    ''',
                    [lms_group_id],
                )
                members = {int(row[0]): row[1] for row in cur.fetchall()}
        except (KeyError, DatabaseError) as error:
            return JsonResponse({"error": "Could not load the LMS group's members.", "details": str(error)}, status=503)
        # The learner the row was created from always gets it, member or not.
        members.setdefault(aptem_id, learner["name"])
        try:
            per_member = []
            for member_id, member_name in sorted(members.items()):
                member_activity_id = f"audit:{uuid.uuid4()}"
                per_member.append((member_id, member_activity_id, _validate_overlay_activity(
                    body.get("activity"), aptem_id=member_id,
                    learner_name=member_name or f"Learner {member_id}",
                    activity_id=member_activity_id,
                )))
        except ValueError as error:
            return JsonResponse({"error": str(error)}, status=400)
        try:
            with transaction.atomic(using=CONN):
                with connections[CONN].cursor() as cur:
                    _ensure_activity_overlay_table(cur)
                    for member_id, member_activity_id, member_payload in per_member:
                        cur.execute(
                            '''
                            insert into "Manual_audit".activity_overrides (aptem_id, activity_id, operation, payload, source_payload, updated_by)
                            values (%s, %s, 'created', %s::jsonb, null, %s)
                            ''',
                            [member_id, member_activity_id, json.dumps(member_payload), updated_by],
                        )
                    cur.execute("select now()")
                    updated_at = cur.fetchone()[0]
        except (KeyError, DatabaseError) as error:
            return JsonResponse({"error": "Could not save the group's activities.", "details": str(error)}, status=503)
        own_id, own_payload = next(
            (member_activity_id, member_payload)
            for member_id, member_activity_id, member_payload in per_member
            if member_id == aptem_id
        )
        return JsonResponse({
            "ok": True, "aptem_id": aptem_id, "activity_id": own_id,
            "operation": "created", "payload": own_payload,
            "group_id": lms_group_id, "created_for": len(per_member),
            "updated_by": updated_by, "updated_at": updated_at.isoformat(),
        })

    try:
        if request.method == "POST":
            activity_id = f"audit:{uuid.uuid4()}"
            payload = _validate_overlay_activity(
                body.get("activity"), aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
            )
            operation = "created"
        elif request.method == "PUT":
            if not activity_id:
                raise ValueError("activity_id is required")
            with connections[CONN].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                cur.execute('''select source_payload from "Manual_audit".activity_overrides where aptem_id = %s and activity_id = %s and operation = 'replaced' ''', [aptem_id, activity_id])
                existing = cur.fetchone()
            raw_source = existing[0] if existing and existing[0] else body.get("snapshot")
            if not isinstance(raw_source, dict):
                raise ValueError("snapshot is required for a reversible date change")
            source_payload = _validate_overlay_activity(
                raw_source, aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
                allow_any_category=True,
            )
            payload = _validate_overlay_activity(
                body.get("activity"), aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
                allow_any_category=True,
            )
            payload["audit_replaced"] = True
            operation = "replaced"
        elif request.method == "PATCH":
            if not activity_id.startswith("audit:"):
                return JsonResponse({"error": "Only audit-created activities can be patched here."}, status=400)
            with connections[CONN].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                cur.execute('''select payload from "Manual_audit".activity_overrides where aptem_id = %s and activity_id = %s and operation = 'created' ''', [aptem_id, activity_id])
                existing = cur.fetchone()
            if not existing:
                return JsonResponse({"error": "Audit activity was not found."}, status=404)
            merged = {**(existing[0] or {}), **(body.get("patch") or {})}
            payload = _validate_overlay_activity(
                merged, aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
                allow_any_category=True,
            )
            operation = "created"
        else:
            if not activity_id:
                raise ValueError("activity_id is required")
            raw_snapshot = body.get("snapshot")
            with connections[CONN].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                cur.execute('''select source_payload from "Manual_audit".activity_overrides where aptem_id = %s and activity_id = %s and operation = 'replaced' ''', [aptem_id, activity_id])
                existing = cur.fetchone()
            if existing and existing[0]:
                raw_snapshot = existing[0]
            if not isinstance(raw_snapshot, dict):
                raise ValueError("snapshot is required for a reversible deletion")
            payload = _validate_overlay_activity(
                raw_snapshot, aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
                allow_any_category=True,
            )
            operation = "deleted"
    except ValueError as error:
        return JsonResponse({"error": str(error)}, status=400)

    try:
        with connections[CONN].cursor() as cur:
            _ensure_activity_overlay_table(cur)
            cur.execute(
                '''
                insert into "Manual_audit".activity_overrides (aptem_id, activity_id, operation, payload, source_payload, updated_by)
                values (%s, %s, %s, %s::jsonb, %s::jsonb, %s)
                on conflict (aptem_id, activity_id) do update set
                    operation = excluded.operation, payload = excluded.payload,
                    source_payload = excluded.source_payload,
                    updated_by = excluded.updated_by, updated_at = now()
                returning updated_at
                ''',
                [aptem_id, activity_id, operation, json.dumps(payload), json.dumps(source_payload) if source_payload else None, updated_by],
            )
            updated_at = cur.fetchone()[0]
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not save activity override.", "details": str(error)}, status=503)
    return JsonResponse({
        "ok": True, "aptem_id": aptem_id, "activity_id": activity_id,
        "operation": operation, "payload": payload,
        "updated_by": updated_by, "updated_at": updated_at.isoformat(),
    })
