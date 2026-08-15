"""Ledger endpoints for the REAL (auditor-copy) workspace.

These power the copy of the Learner Log Pro app (``/workspace/auditor-copy``).
Unlike ``learner_log_views`` — which reads the wide ``Audit.mre`` table — these
views read a single JSON column: ``Audit.learner_match.programme_structure``.

Only the learners on the exact programme ``Level 6 Project Controls Professional``
are served (6 rows at the time of writing). Each of those rows stores a
``programme_structure`` JSON of the shape::

    {
      "programme": "Level 6 Project Controls Professional",
      "months": [
        {
          "month": "October 2024",
          "date": "2024-10-01",
          "hours": {"planned": 24.0, "completed": 51.32},
          "LMS activities": [ {component_id, title, component_kind, material_type,
                               started_at, completed_at, time_spent_seconds,
                               configured_duration_seconds, preview_url, ...}, ... ],
          "attendance": [ {key, date, module, attended, planned_hours, actual_hours}, ... ],
          "assignment": [ {name, status, due_date, completed_date,
                           planned_hours, actual_hours, assignment_id}, ... ]
        }, ...
      ]
    }

The views flatten that per-month shape into the SAME response contract the copy
UI already consumes (see ``learner-log-pro-copy/lib/api.ts``): a learner-summary
list and a paginated activity list. That way the UI components stay unchanged.

Note: ``Audit.learner_match`` lives on the ``enrolment`` Neon database (same as
``learner_api``), so every query here uses ``connections["enrolment"]`` — this
mirrors how ``audit_api/views.py`` reaches the same table.
"""

import collections
import copy
import datetime
import html
import io
import json
import mimetypes
import re
import time
import uuid
from urllib.parse import quote, unquote, urlparse

from django.conf import settings
from django.db import DatabaseError, connections
from django.http import HttpRequest, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET

from .db_source import cache_scope, resolve
from .contract_documents import ensure_contract_archive_table
from .evidence_documents import ensure_evidence_override_table
from .learner_exclusions import is_excluded_learner
from .profile_overrides import apply_break_overrides, apply_profile_overrides, get_profile_overrides

try:
    from azure.storage.blob import BlobSasPermissions, generate_blob_sas
except ImportError:  # pragma: no cover - exercised only when optional Azure SDK is absent.
    BlobSasPermissions = None
    generate_blob_sas = None

try:
    from learner_api.evidence_storage import download_blob_bytes
except ImportError:  # pragma: no cover - optional storage helper in slim test envs.
    download_blob_bytes = None


# The one programme these endpoints serve. Kept as an exact match (not a
# substring/ILIKE) so cohort variants like "... PCP - May 25" are excluded.
PROGRAMME_NAME = "Level 6 Project Controls Professional"

# Audit.learner_match is stored on the enrolment database branch.
CONN = "enrolment"


# --- small parsing helpers -------------------------------------------------

def _to_float(value):
    if value in (None, ""):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _iso_date(value):
    """Pull a ``YYYY-MM-DD`` string out of the many timestamp formats used in
    the JSON (``2024-11-17 19:04:59``, ``2024-10-01T00:00:00+01:00``,
    ``2024-10-01``). Returns None when nothing date-like is present."""
    if not value:
        return None
    match = re.search(r"(\d{4})-(\d{2})-(\d{2})", str(value))
    return match.group(0) if match else None


def _period_of(iso_date):
    return iso_date[:7] if iso_date else None


def _training_plan_from_audit(raw_plan):
    """Normalise the deployed ``Audit.learner_match.aptem_training_plan``."""
    if isinstance(raw_plan, str):
        try:
            raw_plan = json.loads(raw_plan)
        except ValueError:
            raw_plan = []
    if not isinstance(raw_plan, list):
        raw_plan = []

    # Keep a lossless copy as part of the API contract.  The structured fields
    # below drive the current profile UI, while ``raw`` guarantees that a new
    # Aptem field is not silently discarded before the frontend knows how to
    # present it.
    source_plan = copy.deepcopy(raw_plan)
    months = []
    total = 0
    completed = 0
    for month in raw_plan:
        if not isinstance(month, dict):
            continue
        modules = []
        for item in month.get("modules") or []:
            if not isinstance(item, dict):
                continue
            component = item.get("components") if isinstance(item.get("components"), dict) else {}
            status = component.get("status") or "Unknown"
            total += 1
            if str(status).strip().lower() == "completed":
                completed += 1
            modules.append({
                "name": item.get("module") or "Untitled module",
                "type": component.get("type") or "",
                "status": status,
                "components": copy.deepcopy(component),
                "raw": copy.deepcopy(item),
            })
        months.append({
            "month": month.get("month") or "",
            "date": _iso_date(month.get("date")),
            "modules": modules,
            "raw": copy.deepcopy(month),
        })

    return {
        "total_modules": total,
        "completed_modules": completed,
        "months": months,
        "raw": source_plan,
    }


def _clock(value):
    """Pull the ``HH:MM:SS`` clock time out of a timestamp string
    (``2024-11-17 19:04:59``, ``2024-11-17T19:04:59+01:00``), keeping the
    seconds so the ledger shows the exact stamp, not a minute-rounded one. The
    seconds group is optional so a bare ``HH:MM`` value still matches. Anchored
    on the ``T``/space separator so a bare date — or a ``+01:00`` offset — never
    matches. Returns None when there's no time-of-day part."""
    if not value:
        return None
    match = re.search(r"[T ](\d{2}:\d{2}(?::\d{2})?)", str(value))
    return match.group(1) if match else None


def _stamp_parts(value):
    """``(YYYY-MM-DD, HH:MM)`` for a timestamp, or None if either half is
    missing. Ordering these tuples orders the timestamps themselves."""
    date, clock = _iso_date(value), _clock(value)
    return (date, clock) if date and clock else None


def _time_from_to(started_at, completed_at):
    """The ledger's "Timestamp" column: when the learner started and finished,
    e.g. ``19:04 – 19:26``. Falls back to whichever single end we have; None
    when neither timestamp carries a time-of-day."""
    start = _clock(started_at)
    end = _clock(completed_at)
    if start and end:
        return f"{start} – {end}"
    if start:
        return f"{start} – …"
    if end:
        return f"… – {end}"
    return None


def _integer_param(request, name, default, minimum, maximum):
    raw = request.GET.get(name)
    if raw is None or raw == "":
        return default
    value = int(raw)
    if not minimum <= value <= maximum:
        raise ValueError(f"{name} must be between {minimum} and {maximum}")
    return value


def _lms_category(activity):
    """Human-facing category for an LMS activity row. Quizzes have no
    material_type, so fall back to the component kind."""
    material = (activity.get("material_type") or "").strip()
    if material:
        return material
    kind = (activity.get("component_kind") or "").strip()
    return kind or "lesson"


# The full names behind the K/S/B single-letter KSB groups, for display.
_KSB_TYPE_LABELS = {"K": "Knowledge", "S": "Skill", "B": "Behaviour"}


def _completion_records(item):
    """Clean an activity's per-completion history (each start/finish + time)."""
    records = []
    for record in item.get("completion_records") or []:
        if not isinstance(record, dict):
            continue
        records.append({
            "record_id": record.get("record_id"),
            "started_at": record.get("started_at"),
            "completed_at": record.get("completed_at"),
            "time_spent_seconds": record.get("time_spent_seconds"),
            "time_spent_formatted": record.get("time_spent_formatted"),
        })
    return records


def _normalize_ksbs(ksbs):
    """Flatten the ``{"K": [...], "S": [...], "B": [...]}`` KSB block that now
    lives on every activity in programme_structure into a single ordered list of
    ``{code, type, type_label, description, reason}`` dicts (Knowledge, then
    Skill, then Behaviour). Returns ``[]`` when no KSBs are present."""
    if not isinstance(ksbs, dict):
        return []
    result = []
    for group in ("K", "S", "B"):
        items = ksbs.get(group)
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            code = (item.get("code") or "").strip()
            if not code:
                continue
            group_letter = (item.get("type") or group or "").strip().upper()[:1] or group
            result.append({
                "code": code,
                "type": group_letter,
                "type_label": _KSB_TYPE_LABELS.get(group_letter, group_letter),
                "description": (item.get("description") or "").strip() or None,
                "reason": (item.get("reason") or "").strip() or None,
            })
    return result


# --- attendance-session grouping + recording matching ----------------------
#
# An attendance key is ``<learner_id>_<YYYY-MM-DD>_<group>`` — e.g.
# ``471_2026-07-03_project_planning_and_control_ppc_andrew``. Two learners
# attended the *same* live session when the last two parts (date + group) match;
# the group text is spelled inconsistently across learners (``ray-project...``
# vs ``ray_project...`` vs ``Ray - Project ...``) so it's normalised first.
#
# The recording of that session lives as ordinary LMS lesson(s) whose *title*
# starts with the session date in ``D-M-YYYY`` form and names the course, e.g.
# ``26-6-2026 Project Planning & Control - Andrew Millington Part 1``. So a
# recording belongs to a session when its title-date equals the session date and
# its course name overlaps the session group. (Per-activity ``created_at`` /
# ``completed_at`` are upload / watch timestamps and don't identify the session.)

_ATT_KEY_RE = re.compile(r"^(\d+)_(\d{4}-\d{2}-\d{2})_(.+)$")
# A D-M-YYYY / D/M/YYYY date appearing ANYWHERE in a recording title. Some
# courses date-stamp the front ("26-6-2026 Project Planning…"), others the tail
# ("… Customer Experience - 17/7/2026"), so it is not start-anchored.
_TITLE_DATE_RE = re.compile(r"(?<!\d)(\d{1,2})[-/](\d{1,2})[-/](\d{4})(?!\d)")

# Filler words that carry no course identity — dropped before comparing a
# session group against a recording's course name.
_TOKEN_STOP = {
    "and", "the", "of", "a", "an", "part", "parts", "session", "sessions",
    "recorded", "live", "lecture", "lesson", "qualification", "module",
    "group", "online",
}


def _norm_group(value):
    """Lower-case and collapse all non-alphanumerics to single spaces, so the
    ``-``/``_``/space/case spelling variants of a group name unify."""
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _significant_tokens(text):
    """Meaningful word set of a group/course name (stop-words and bare numbers
    removed) — used to score whether a recording's course matches a session."""
    return {t for t in _norm_group(text).split() if t not in _TOKEN_STOP and not t.isdigit()}


def _parse_attendance_key(key):
    match = _ATT_KEY_RE.match(key or "")
    if not match:
        return None
    return match.group(1), match.group(2), match.group(3)


def _clean_title(title):
    """Un-escape HTML entities (``&amp;``) and drop zero-width spaces that some
    recording titles carry, so parsing/matching sees plain text."""
    return html.unescape(title or "").replace("​", " ")


def _title_session_date(title):
    """Parse the ``D-M-YYYY`` date somewhere in a recording's title into an ISO
    ``YYYY-MM-DD`` string (the live-session date), or None if absent."""
    match = _TITLE_DATE_RE.search(_clean_title(title))
    if not match:
        return None
    day, month, year = (int(part) for part in match.groups())
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return None
    return f"{year:04d}-{month:02d}-{day:02d}"


def _recording_course(title):
    """Strip the date and any ``Part N`` marker out of a recording title,
    leaving just the course/topic words for token matching."""
    text = _TITLE_DATE_RE.sub(" ", _clean_title(title))
    text = re.sub(r"(?i)\bpart\s*\d+\b", " ", text)
    return text.strip(" -–—")


# --- shape mapping ---------------------------------------------------------

def _reading_quiz_rows(aptem_id, learner_name, month_no, month_unit, month_date,
                       planned_month_period, wrapper):
    """Turn a ``kind: "reading+quiz"`` week-group wrapper into ONE bundle activity
    row. The wrapper's nested materials (readings, quizzes, pdfs) are returned as
    a ``components`` list on that single row — so the group shows as one entry in
    the ledger and its detailed contents are revealed on the activity page. The
    group's ``week`` and ``KSBs`` describe the whole bundle."""
    items = [item for item in (wrapper.get("items") or []) if isinstance(item, dict)]
    if not items:
        return []
    group_week = wrapper.get("week") or None
    group_ksbs = _normalize_ksbs(wrapper.get("KSBs"))

    components = []
    planned_seconds_total = 0.0
    actual_seconds_total = 0.0
    completed_count = 0
    last_date = None
    # Earliest start / latest finish across the bundle's materials, each kept as
    # a (date, HH:MM) pair so the span can be rejected when it crosses days.
    first_start = None
    last_end = None
    category_counts = collections.Counter()
    for item in items:
        act = item.get("activity") if isinstance(item.get("activity"), dict) else {}
        duration = act.get("content_duration") if isinstance(act.get("content_duration"), dict) else {}
        planned_seconds = duration.get("seconds")
        if isinstance(planned_seconds, (int, float)):
            planned_seconds_total += planned_seconds
        # Actual = engineered OTJH-credited seconds (activity["otjh"]), not raw time.
        item_otjh = act.get("otjh") if isinstance(act.get("otjh"), dict) else {}
        otjh_seconds = _to_float(item_otjh.get("seconds"))
        spent = otjh_seconds if otjh_seconds is not None else act.get("time_spent_seconds")
        if isinstance(spent, (int, float)):
            actual_seconds_total += spent
        done = bool(act.get("completed"))
        if done:
            completed_count += 1
        item_date = _iso_date(act.get("completed_at")) or _iso_date(act.get("started_at"))
        if item_date and (last_date is None or item_date > last_date):
            last_date = item_date
        started = _stamp_parts(act.get("started_at"))
        if started and (first_start is None or started < first_start):
            first_start = started
        ended = _stamp_parts(act.get("completed_at"))
        if ended and (last_end is None or ended > last_end):
            last_end = ended
        material = (item.get("material_type") or "").strip() or "reading"
        category_counts[material] += 1
        attempt_records = act.get("quiz_attempt_records") or 0
        components.append({
            "component_id": item.get("component_id"),
            "title": item.get("title") or "Untitled material",
            "material_type": material,
            "material_format": act.get("material_format") or None,
            "iframe_url": item.get("iframe_url") or None,
            "status": act.get("status") or None,
            "done": done,
            "activity_date": item_date,
            "planned_hours": round(planned_seconds / 3600, 2) if isinstance(planned_seconds, (int, float)) and planned_seconds else None,
            "time_spent_formatted": item_otjh.get("logged_time") or act.get("time_spent_formatted") or None,
            "otjh_hours": _to_float(item_otjh.get("hours")),
            "otjh_credited": bool(item_otjh.get("credited")) if item_otjh else None,
            "attempt_number": act.get("attempt_number"),
            "highest_score": act.get("highest_score"),
            # Quiz result fields (null for readings/pdfs).
            "score_percent": act.get("score_percent"),
            "answered_questions": act.get("answered_questions"),
            "correct_answers": act.get("correct_answers"),
            "incorrect_answers": act.get("incorrect_answers"),
            # Whether a graded body exists to open (from the quiz_attempts column).
            "has_body": material == "quiz" and bool(attempt_records),
        })

    total = len(components)
    month_period = _period_of(last_date) or planned_month_period
    # A bundle spans several materials, so its span is only meaningful when they
    # were all worked on the same day — otherwise "19:04 – 21:30" would silently
    # hide a multi-day gap.
    same_day_span = (
        (first_start[1], last_end[1])
        if first_start and last_end and first_start[0] == last_end[0]
        else None
    )
    # Bundle id is stable across learners for the same week group, so the
    # cross-learner "who has this" view lines them up.
    week_slug = _norm_group(group_week).replace(" ", "-") if group_week else f"m{month_no}"
    bundle_id = f"rqb:{week_slug}"
    # A concise "quizzes / readings / pdfs" descriptor for the ledger row.
    breakdown = ", ".join(f"{count} {name}" for name, count in category_counts.most_common())
    return [{
        "id": f"{aptem_id}:{month_no}:rqb:{week_slug}",
        "mre_id": bundle_id,
        "learner": learner_name,
        "plan_id": bundle_id,
        "month_no": month_no,
        "month_unit": month_unit,
        "unit_planned_date": month_date,
        "activity_date": last_date,
        "learner_activity_date": last_date,
        "activity_period": month_period,
        "time_from_to": (
            f"{same_day_span[0]} – {same_day_span[1]}" if same_day_span else None
        ),
        "time_from": same_day_span[0] if same_day_span else None,
        "time_to": same_day_span[1] if same_day_span else None,
        "actual_lms_hours": round(actual_seconds_total / 3600, 2) if actual_seconds_total else None,
        "activity_category": "reading+quiz",
        "activity_unit": group_week or "Readings & quizzes",
        "section_title": group_week or None,
        # Completion counts are intentionally omitted for now — just the make-up.
        "activity_description": breakdown or None,
        "delivery_method": f"Readings & quizzes ({total} items)",
        "planned_hours": round(planned_seconds_total / 3600, 2) if planned_seconds_total else None,
        "source_course": None,
        "source_url": None,
        "source_basis": "reading+quiz",
        "created_at": None,
        "configured_duration": None,
        "week": group_week,
        "ksbs": group_ksbs,
        "components": components,
        "completed_count": completed_count,
        "component_total": total,
        "done": total > 0 and completed_count == total,
    }]


def _activities_for_row(aptem_id, learner_name, structure):
    """Flatten one learner's programme_structure into activity dicts matching
    the copy UI's ``LearnerActivity`` type (minus the removed KSB / evidence /
    week_sequence fields)."""
    activities = []
    months = structure.get("months") if isinstance(structure, dict) else None
    if not isinstance(months, list):
        return activities

    for month_index, month in enumerate(months):
        if not isinstance(month, dict):
            continue
        month_no = month_index + 1
        month_unit = month.get("month") or ""
        month_date = _iso_date(month.get("date"))
        planned_month_period = _period_of(month_date)

        # 1) LMS activities (videos, pdfs, quizzes, live lessons, ...)
        for index, item in enumerate(month.get("LMS activities") or []):
            if not isinstance(item, dict):
                continue
            # Week-group wrapper (shape: kind/week/items/_meta, no component_id at
            # the top level). "reading+quiz" groups nest their real materials under
            # `items`, each with its own per-learner `activity` — expand those into
            # rows instead of dropping the whole wrapper.
            if item.get("component_id") is None:
                if isinstance(item.get("items"), list):
                    activities.extend(_reading_quiz_rows(
                        aptem_id, learner_name, month_no, month_unit, month_date,
                        planned_month_period, item,
                    ))
                continue
            component_id = item.get("component_id")
            activity_date = _iso_date(item.get("completed_at")) or _iso_date(item.get("started_at"))
            # An activity's "period" is the month it actually happened; if it was
            # never started, bucket it under its planned month.
            month_period = _period_of(activity_date) or planned_month_period
            # Actual = the engineered OTJH-credited hours for this activity
            # (item["otjh"]["hours"]); fall back to raw tracked time when absent.
            otjh = item.get("otjh") if isinstance(item.get("otjh"), dict) else {}
            otjh_hours = _to_float(otjh.get("hours"))
            seconds = item.get("time_spent_seconds")
            actual_hours = otjh_hours if otjh_hours is not None else (
                round(seconds / 3600, 2) if isinstance(seconds, (int, float)) and seconds else None
            )
            planned_seconds = item.get("configured_duration_seconds")
            planned_hours = round(planned_seconds / 3600, 2) if isinstance(planned_seconds, (int, float)) and planned_seconds else None
            activities.append({
                "id": f"{aptem_id}:{month_no}:lms:{component_id}:{index}",
                "mre_id": str(component_id) if component_id is not None else "",
                "learner": learner_name,
                "plan_id": str(component_id) if component_id is not None else "",
                "month_no": month_no,
                "month_unit": month_unit,
                "unit_planned_date": month_date,
                "activity_date": activity_date,
                "learner_activity_date": activity_date,
                "activity_period": month_period,
                "time_from_to": _time_from_to(item.get("started_at"), item.get("completed_at")),
                # Same span split out, for the two-column "From / To" header.
                "time_from": _clock(item.get("started_at")),
                "time_to": _clock(item.get("completed_at")),
                "actual_lms_hours": actual_hours,
                "otjh_credited": bool(otjh.get("credited")) if otjh else None,
                "activity_category": _lms_category(item),
                "activity_unit": item.get("title") or "Untitled activity",
                "section_title": item.get("section_title") or item.get("section") or item.get("week") or None,
                "activity_description": None,
                "delivery_method": item.get("component_kind") or "",
                "planned_hours": planned_hours,
                "source_course": None,
                "source_url": item.get("preview_url") or None,
                "source_basis": item.get("post_type") or None,
                "created_at": item.get("created_at") or None,
                "configured_duration": item.get("configured_duration") or None,
                "week": item.get("week") or None,
                "ksbs": _normalize_ksbs(item.get("KSBs")),
                "completion_records": _completion_records(item),
                "done": bool(item.get("completed")),
            })

        # 2) Assignments
        for index, item in enumerate(month.get("assignment") or []):
            if not isinstance(item, dict):
                continue
            # Real completion date drives "when it happened"; a due date only
            # helps place an un-submitted assignment into a month bucket.
            activity_date = _iso_date(item.get("completed_date"))
            month_period = _period_of(activity_date) or _period_of(item.get("due_date")) or planned_month_period
            activities.append({
                "id": f"{aptem_id}:{month_no}:asg:{item.get('assignment_id')}:{index}",
                "mre_id": str(item.get("assignment_id") or ""),
                "learner": learner_name,
                "plan_id": str(item.get("assignment_id") or ""),
                "month_no": month_no,
                "month_unit": month_unit,
                "unit_planned_date": month_date,
                "activity_date": activity_date,
                "learner_activity_date": activity_date,
                "activity_period": month_period,
                # Assignments/attendance carry dates only — no time of day.
                "time_from_to": None,
                "time_from": None,
                "time_to": None,
                "actual_lms_hours": _to_float(item.get("actual_hours")),
                "activity_category": "assignment",
                "activity_unit": item.get("name") or "Assignment",
                "section_title": item.get("section_title") or item.get("section") or None,
                "activity_description": item.get("status") or None,
                "delivery_method": "assignment",
                "planned_hours": _to_float(item.get("planned_hours")),
                "source_course": None,
                "source_url": None,
                "source_basis": "assignment",
                "created_at": None,
                "configured_duration": None,
                "week": None,
                "ksbs": _normalize_ksbs(item.get("KSBs")),
                "done": (item.get("status") or "").lower() == "completed",
            })

        # 3) Attendance (live sessions)
        for index, item in enumerate(month.get("attendance") or []):
            if not isinstance(item, dict):
                continue
            activity_date = _iso_date(item.get("date"))
            month_period = _period_of(activity_date) or planned_month_period
            attended = item.get("attended")
            # Session identity from the key (date + normalised group) — lets us
            # gather every learner who attended the same live session.
            parsed_key = _parse_attendance_key(item.get("key"))
            session_date = parsed_key[1] if parsed_key else activity_date
            session_group = _norm_group(parsed_key[2]) if parsed_key else ""
            activities.append({
                "id": f"{aptem_id}:{month_no}:att:{item.get('key')}:{index}",
                "mre_id": str(item.get("key") or ""),
                "learner": learner_name,
                "plan_id": str(item.get("key") or ""),
                "month_no": month_no,
                "month_unit": month_unit,
                "unit_planned_date": month_date,
                "activity_date": activity_date,
                "learner_activity_date": activity_date,
                "activity_period": month_period,
                # Assignments/attendance carry dates only — no time of day.
                "time_from_to": None,
                "time_from": None,
                "time_to": None,
                "actual_lms_hours": _to_float(item.get("actual_hours")),
                "activity_category": "attendance",
                "activity_unit": item.get("module") or "Attendance",
                "section_title": item.get("section_title") or item.get("section") or None,
                "activity_description": None,
                "delivery_method": "attendance",
                "planned_hours": _to_float(item.get("planned_hours")),
                "source_course": None,
                "source_url": None,
                "source_basis": "attendance",
                "created_at": None,
                "configured_duration": None,
                "week": None,
                "ksbs": _normalize_ksbs(item.get("KSBs")),
                "session_date": session_date,
                "session_group": session_group,
                "done": bool(attended),
            })

    return activities


def _month_hours_for_row(structure):
    """Map each month's curated ``hours`` block to its period. These are the
    authoritative planned/completed hour totals (the per-activity ``time_spent``
    tracking in the LMS is noisy, so summaries use these instead)."""
    result = {}
    months = structure.get("months") if isinstance(structure, dict) else None
    if not isinstance(months, list):
        return result
    for month in months:
        if not isinstance(month, dict):
            continue
        period = _period_of(_iso_date(month.get("date")))
        hours = month.get("hours") if isinstance(month.get("hours"), dict) else {}
        if not period:
            continue
        result[period] = {
            "planned": _to_float(hours.get("planned")) or 0.0,
            "completed": _to_float(hours.get("completed")) or 0.0,
        }
    return result


def _otjh_for_row(structure):
    """Pull the off-the-job-hours adjustment data out of a learner's structure:
    the top-level ``otjh_summary`` and a ``{period -> months[].otjh_adjustment}``
    map. These carry the engineered per-month hour breakdown (att/asg/lms vs the
    Aptem actual) and the review ``flagged`` status."""
    summary = structure.get("otjh_summary") if isinstance(structure, dict) else None
    summary = summary if isinstance(summary, dict) else None
    months_map = {}
    months = structure.get("months") if isinstance(structure, dict) else None
    if isinstance(months, list):
        for month in months:
            if not isinstance(month, dict):
                continue
            adjustment = month.get("otjh_adjustment")
            if not isinstance(adjustment, dict):
                continue
            period = _period_of(_iso_date(month.get("date")))
            if period:
                months_map[period] = adjustment
    return summary, months_map


# Each programme_structure blob is multi-megabyte (thousands of activities), and
# the UI fires several calls per page. Cache the flattened result briefly so the
# repeated fetch+parse+flatten cost is paid once, not on every request.
_CACHE_TTL_SECONDS = 20
# Keyed by database source: the live workspace and the HOURS-TEST clone read
# different rows, so they must never serve each other's cached copy.
_cache = {}


def _load_rows():
    now = time.monotonic()
    entry = _cache.get(cache_scope())
    if entry and entry["rows"] is not None and now < entry["expires_at"]:
        return entry["rows"]
    rows = _fetch_rows()
    _cache[cache_scope()] = {"rows": rows, "expires_at": now + _CACHE_TTL_SECONDS}
    return rows


def _fetch_rows():
    """Fetch the 6 exact-programme learners from Audit.learner_match and return
    one dict per learner: aptem_id, name, flattened activities, and the curated
    per-period hour totals."""
    with connections[resolve(CONN)].cursor() as cur:
        # The column is `json` (not jsonb). Using `->>` directly (no ::jsonb
        # cast) avoids converting every row's multi-MB blob to jsonb just to
        # read the programme name — roughly halves the query time.
        cur.execute(
            '''
            select lm.aptem_id, lm.learner_name, lm.learner_email,
                   lm.programme_structure, lm.aptem_training_plan,
                   lm.programme_name, profile.program_status,
                   profile.break_in_learning, owner.coach_name,
                   owner.coach_email
            from "Audit".learner_match lm
            left join lateral (
                select program_status, "Break in learning" as break_in_learning
                from fetching_evidence.aptem_cv_contracts_probe
                where learner_id = lm.aptem_id
                order by fetched_at desc nulls last, id desc
                limit 1
            ) profile on true
            left join lateral (
                select "OwnerName" as coach_name, "OwnerEmail" as coach_email
                from "LMS"."Aptem_users"
                where "ID" = lm.aptem_id
                limit 1
            ) owner on true
            where lm.programme_structure ->> 'programme' = %s
            order by lm.learner_name, lm.aptem_id
            ''',
            [PROGRAMME_NAME],
        )
        rows = cur.fetchall()

    learners = []
    for (
        aptem_id, learner_name, learner_email, structure, training_plan,
        programme_name, program_status, break_in_learning, coach_name,
        coach_email,
    ) in rows:
        if is_excluded_learner(aptem_id, learner_name):
            continue
        # psycopg may hand json/jsonb back as a decoded dict or as raw text
        # depending on the configured loaders — normalise to a dict either way.
        if isinstance(structure, str):
            try:
                structure = json.loads(structure)
            except ValueError:
                structure = None
        if isinstance(training_plan, str):
            try:
                training_plan = json.loads(training_plan)
            except ValueError:
                training_plan = None
        if isinstance(break_in_learning, str):
            try:
                break_in_learning = json.loads(break_in_learning)
            except ValueError:
                break_in_learning = None
        break_in_learning = break_in_learning if isinstance(break_in_learning, dict) else {}
        name = learner_name or f"Learner {aptem_id}"
        otjh_summary, otjh_months = _otjh_for_row(structure)
        learners.append({
            "aptem_id": aptem_id,
            "name": name,
            "otjh_summary": otjh_summary,
            "otjh_months": otjh_months,
            "email": learner_email,
            "programme_name": programme_name or PROGRAMME_NAME,
            "program_status": program_status or "Unknown",
            "has_break_in_learning": (
                bool(break_in_learning.get("has_break_in_learning"))
                or str(program_status or "").strip().lower() == "onbreak"
            ),
            "coach": {
                "name": coach_name or None,
                "email": coach_email or None,
            },
            "training_plan": training_plan if isinstance(training_plan, list) else [],
            "activities": _activities_for_row(aptem_id, name, structure),
            "month_hours": _month_hours_for_row(structure),
        })
    return learners


def _fetch_profile_row(learner_key):
    """Fetch one profile row across all programmes.

    The cached ledger rows are intentionally scoped to the original PCP summary
    cohort. Profile links now come from the live all-programme cohort feed, so a
    learner outside that subset still needs the same rich profile payload.
    """
    with connections[resolve(CONN)].cursor() as cur:
        cur.execute(
            '''
            select lm.aptem_id, lm.learner_name, lm.learner_email,
                   lm.programme_structure, lm.aptem_training_plan,
                   lm.programme_name, profile.program_status,
                   profile.break_in_learning, owner.coach_name,
                   owner.coach_email
            from "Audit".learner_match lm
            left join lateral (
                select program_status, "Break in learning" as break_in_learning
                from fetching_evidence.aptem_cv_contracts_probe
                where learner_id = lm.aptem_id
                order by fetched_at desc nulls last, id desc
                limit 1
            ) profile on true
            left join lateral (
                select "OwnerName" as coach_name, "OwnerEmail" as coach_email
                from "LMS"."Aptem_users"
                where "ID" = lm.aptem_id
                limit 1
            ) owner on true
            where lm.aptem_id::text = %s
               or lower(btrim(lm.learner_name)) = %s
            order by lm.learner_name, lm.aptem_id
            limit 1
            ''',
            [learner_key, learner_key],
        )
        row = cur.fetchone()
    if row is None:
        return _fetch_profile_source_row(learner_key)

    (
        aptem_id, learner_name, learner_email, structure, training_plan,
        programme_name, program_status, break_in_learning, coach_name,
        coach_email,
    ) = row
    if is_excluded_learner(aptem_id, learner_name):
        return None
    if isinstance(structure, str):
        try:
            structure = json.loads(structure)
        except ValueError:
            structure = None
    if isinstance(training_plan, str):
        try:
            training_plan = json.loads(training_plan)
        except ValueError:
            training_plan = None
    if isinstance(break_in_learning, str):
        try:
            break_in_learning = json.loads(break_in_learning)
        except ValueError:
            break_in_learning = None
    break_in_learning = break_in_learning if isinstance(break_in_learning, dict) else {}
    name = learner_name or f"Learner {aptem_id}"
    otjh_summary, otjh_months = _otjh_for_row(structure)
    return {
        "aptem_id": aptem_id,
        "name": name,
        "otjh_summary": otjh_summary,
        "otjh_months": otjh_months,
        "email": learner_email,
        "programme_name": programme_name or PROGRAMME_NAME,
        "program_status": program_status or "Unknown",
        "has_break_in_learning": (
            bool(break_in_learning.get("has_break_in_learning"))
            or str(program_status or "").strip().lower() == "onbreak"
        ),
        "coach": {
            "name": coach_name or None,
            "email": coach_email or None,
        },
        "training_plan": training_plan if isinstance(training_plan, list) else [],
        "activities": _activities_for_row(aptem_id, name, structure),
        "month_hours": _month_hours_for_row(structure),
    }


def _fetch_profile_source_row(learner_key):
    """Build a profile shell from fetching_evidence when learner_match lags.

    Newly added learners can appear in the live cohort/contracts data before the
    nightly learner_match enrichment has produced programme_structure. Returning
    a same-shaped learner dict lets the profile page auto-exist immediately; the
    rich sections that are already keyed by aptem_id are still loaded below by
    _load_profile_sources.
    """
    with connections[resolve(CONN)].cursor() as cur:
        cur.execute(
            '''
            select contracts.learner_id, contracts.full_name, contracts.email,
                   contracts.program_name, contracts.program_status,
                   contracts."Break in learning", owner.coach_name,
                   owner.coach_email
            from fetching_evidence.aptem_cv_contracts_probe contracts
            left join lateral (
                select "OwnerName" as coach_name, "OwnerEmail" as coach_email
                from "LMS"."Aptem_users"
                where "ID" = contracts.learner_id
                limit 1
            ) owner on true
            where contracts.learner_id::text = %s
               or lower(btrim(contracts.full_name)) = %s
               or lower(btrim(contracts.email)) = %s
            order by contracts.fetched_at desc nulls last, contracts.id desc
            limit 1
            ''',
            [learner_key, learner_key, learner_key],
        )
        row = cur.fetchone()
    if row is None or row[0] is None:
        return None

    aptem_id, learner_name, learner_email, programme_name, program_status, break_in_learning, coach_name, coach_email = row
    if is_excluded_learner(aptem_id, learner_name):
        return None
    if isinstance(break_in_learning, str):
        try:
            break_in_learning = json.loads(break_in_learning)
        except ValueError:
            break_in_learning = None
    break_in_learning = break_in_learning if isinstance(break_in_learning, dict) else {}
    name = learner_name or f"Learner {aptem_id}"
    return {
        "aptem_id": aptem_id,
        "name": name,
        "otjh_summary": {},
        "otjh_months": {},
        "email": learner_email,
        "programme_name": programme_name or "Unknown programme",
        "program_status": program_status or "Unknown",
        "has_break_in_learning": (
            bool(break_in_learning.get("has_break_in_learning"))
            or str(program_status or "").strip().lower() == "onbreak"
        ),
        "coach": {
            "name": coach_name or None,
            "email": coach_email or None,
        },
        "training_plan": [],
        "activities": [],
        "month_hours": {},
    }


def _learner_id(name):
    """Stable slug used as the learner filter key. Mirrors the original MRE
    app convention (id == lowercased name) so the existing UI links/dropdowns
    keep working without change."""
    return (name or "").strip().lower()


def _contract_blob_from_azure_path(value):
    """Parse `az://account/container/blob` contract references."""
    text = str(value or "").strip()
    if not text:
        return None
    parsed = urlparse(text)
    if parsed.scheme == "az":
        account = (parsed.netloc or "").strip()
        path = parsed.path.lstrip("/")
        if "/" not in path:
            return None
        container, blob_name = path.split("/", 1)
        configured_account = getattr(settings, "AZURE_STORAGE_ACCOUNT", "")
        if configured_account and account and account.lower() != configured_account.lower():
            return None
        return container, unquote(blob_name)
    if parsed.scheme == "https" and parsed.hostname:
        configured_account = getattr(settings, "AZURE_STORAGE_ACCOUNT", "")
        expected_host = f"{configured_account}.blob.core.windows.net".lower()
        if configured_account and parsed.hostname.lower() != expected_host:
            return None
        path = parsed.path.lstrip("/")
        if "/" not in path:
            return None
        container, blob_name = path.split("/", 1)
        return container, unquote(blob_name)
    return None


def _safe_contract_filename(document_name, blob_name):
    blob_filename = unquote(blob_name.rsplit("/", 1)[-1]) if blob_name else ""
    filename = blob_filename or str(document_name or "contract").strip() or "contract"
    filename = re.sub(r'[\r\n"\\]+', " ", filename).strip()
    return filename or "contract"


def _contract_preview_url(azure_path, document_name):
    location = _contract_blob_from_azure_path(azure_path)
    if not location:
        return ""
    account_name = getattr(settings, "AZURE_STORAGE_ACCOUNT", "")
    account_key = getattr(settings, "AZURE_STORAGE_KEY", "")
    if not account_name or not account_key or generate_blob_sas is None or BlobSasPermissions is None:
        return ""

    container, blob_name = location
    filename = _safe_contract_filename(document_name, blob_name)
    content_type = mimetypes.guess_type(filename)[0] or mimetypes.guess_type(blob_name)[0] or "application/octet-stream"
    token = generate_blob_sas(
        account_name=account_name,
        container_name=container,
        blob_name=blob_name,
        account_key=account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(minutes=30),
        content_type=content_type,
        content_disposition=f'inline; filename="{filename}"',
    )
    return f"https://{account_name}.blob.core.windows.net/{container}/{quote(blob_name, safe='/')}?{token}"


_CONTRACT_SIGNATURE_CACHE_TTL_SECONDS = 600
_contract_signature_cache = {}


def _normalise_contract_date(value):
    text = str(value or "").strip()
    if not text:
        return None
    iso_match = re.search(r"(?<!\d)(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})(?!\d)", text)
    if iso_match:
        year, month, day = (int(part) for part in iso_match.groups())
    else:
        uk_match = re.search(r"(?<!\d)(\d{1,2})[-/.](\d{1,2})[-/.](\d{2,4})(?!\d)", text)
        if not uk_match:
            return None
        day, month, year = (int(part) for part in uk_match.groups())
        if year < 100:
            year += 2000
    try:
        return datetime.date(year, month, day).isoformat()
    except ValueError:
        return None


def _extract_contract_text(filename, data):
    extension = "." + str(filename or "").rsplit(".", 1)[-1].lower() if "." in str(filename or "") else ""
    if extension == ".pdf":
        try:
            import fitz

            document = fitz.open(stream=data, filetype="pdf")
            lines = []
            for page_index in range(min(len(document), 30)):
                page = document[page_index]
                text = re.sub(r"\s+", " ", page.get_text("text") or "").strip()
                if text:
                    lines.append(f"Page {page_index + 1}: {text}")
            text = "\n".join(lines)
            if text:
                return text
        except Exception:
            pass
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            lines = []
            for page in reader.pages[:10]:
                text = re.sub(r"\s+", " ", page.extract_text() or "").strip()
                if text:
                    lines.append(text)
            return "\n".join(lines)
        except Exception:
            return ""
    if extension == ".docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            lines = []
            for paragraph in document.paragraphs:
                text = re.sub(r"\s+", " ", paragraph.text).strip()
                if text:
                    lines.append(text)
            for table in document.tables:
                for row in table.rows:
                    values = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        lines.append(" | ".join(values))
            return "\n".join(lines)
        except Exception:
            return ""
    if extension in {".txt", ".csv"}:
        return data.decode("utf-8-sig", errors="replace")
    return ""


def _contract_signature_dates_from_text(text):
    clean = re.sub(r"\s+", " ", str(text or "")).strip()
    if not clean:
        return {"learner_signed_date": None, "fully_signed_date": None}

    date_pattern = r"(\d{4}[-/.]\d{1,2}[-/.]\d{1,2}|\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})"
    signature_section = clean[-3500:]
    markers = []
    for marker_pattern in (
        r"\bsignatures?\s*&\s*declarations?\b",
        r"\bsignatories\b",
        r"\bsignatures\b",
    ):
        markers.extend(re.finditer(marker_pattern, clean, flags=re.IGNORECASE))
    if markers:
        marker = max(markers, key=lambda item: item.start())
        signature_section = clean[marker.start(): marker.start() + 3500]

    learner_date = None
    # Prefer the date in the apprentice/learner signatory row.
    for pattern in (
        rf"(?:apprentice|learner)\s*:?.{{0,220}}?\bdate\s*:?\s*{date_pattern}",
        rf"(?:apprentice|learner)\s*:?.{{0,220}}?{date_pattern}",
    ):
        match = re.search(pattern, signature_section, flags=re.IGNORECASE)
        if match:
            learner_date = _normalise_contract_date(match.group(1))
            if learner_date:
                break

    all_dates = [
        parsed for parsed in (_normalise_contract_date(match.group(1)) for match in re.finditer(date_pattern, signature_section))
        if parsed
    ]
    return {
        "learner_signed_date": learner_date,
        "fully_signed_date": max(all_dates) if all_dates else None,
    }


def _contract_signature_dates(azure_path, document_name):
    location = _contract_blob_from_azure_path(azure_path)
    if not location or download_blob_bytes is None:
        return {"learner_signed_date": None, "fully_signed_date": None}
    now = time.monotonic()
    cached = _contract_signature_cache.get(str(azure_path))
    if cached and now < cached["expires_at"]:
        return cached["value"]

    result = {"learner_signed_date": None, "fully_signed_date": None}
    try:
        container, blob_name = location
        filename = _safe_contract_filename(document_name, blob_name)
        data = download_blob_bytes(container, blob_name)
        text = _extract_contract_text(filename, data)
        result = _contract_signature_dates_from_text(text)
    except Exception:
        result = {"learner_signed_date": None, "fully_signed_date": None}

    _contract_signature_cache[str(azure_path)] = {
        "expires_at": now + _CONTRACT_SIGNATURE_CACHE_TTL_SECONDS,
        "value": result,
    }
    return result


# --- views -----------------------------------------------------------------

@require_GET
def health(_request: HttpRequest) -> JsonResponse:
    alias = resolve(CONN)
    with connections[alias].cursor() as cursor:
        cursor.execute("SELECT current_database(), now()")
        database, timestamp = cursor.fetchone()
    return JsonResponse({"ok": True, "database": database, "time": timestamp})


@require_GET
def learner_summaries(request: HttpRequest) -> JsonResponse:
    """Per-learner totals + the filter option lists (learners/periods/
    categories/months) the copy UI needs. Matches ``getLearners()``."""
    try:
        period = request.GET.get("period", "").strip()
        search = request.GET.get("search", "").strip().lower()[:100]
        position = request.GET.get("position", "").strip().lower()
        if period and (len(period) != 7 or period[4] != "-" or not period.replace("-", "").isdigit()):
            raise ValueError("period must use YYYY-MM format")
        if position not in ("", "behind", "ahead"):
            raise ValueError("position must be behind or ahead")
    except (TypeError, ValueError) as error:
        return JsonResponse({"error": "Invalid query parameters", "details": str(error)}, status=400)

    try:
        learners_raw = _load_rows()
    except (KeyError, DatabaseError) as error:
        return JsonResponse(
            {"error": "Could not read Audit.learner_match from the enrolment database.", "details": str(error)},
            status=503,
        )

    # Only surface months up to the current one, and only those where real
    # activity happened — so the dropdown never defaults to empty future
    # scaffolding (the structure is planned out to 2028).
    current_period = datetime.date.today().strftime("%Y-%m")
    period_values = set()
    categories = set()
    month_labels = {}
    for learner in learners_raw:
        for activity in learner["activities"]:
            if activity["activity_category"]:
                categories.add(activity["activity_category"])
            bucket = activity["activity_period"]
            if bucket and activity["activity_date"] and bucket <= current_period:
                period_values.add(bucket)
                month_labels[bucket] = activity["month_unit"] or bucket

    summaries = []
    for learner in learners_raw:
        month_hours = learner["month_hours"]
        if period:
            entry_hours = month_hours.get(period, {"planned": 0.0, "completed": 0.0})
            planned = entry_hours["planned"]
            actual = entry_hours["completed"]
            activities = [item for item in learner["activities"] if item["activity_period"] == period]
        else:
            # Cumulative programme totals (all planned months + all completed).
            planned = sum(value["planned"] for value in month_hours.values())
            actual = sum(value["completed"] for value in month_hours.values())
            activities = learner["activities"]
        dates = [
            item["activity_date"] for item in activities
            if item["activity_date"] and item["activity_date"][:7] <= current_period
        ]
        otjh_summary = learner.get("otjh_summary") or {}
        otjh_months = learner.get("otjh_months") or {}
        month_adjustment = otjh_months.get(period) if period else None
        summaries.append({
            "id": _learner_id(learner["name"]),
            "name": learner["name"],
            "entries": len(activities),
            "planned_hours": round(planned, 2),
            "actual_hours": round(actual, 2),
            "gap_hours": round(actual - planned, 2),
            "last_activity_date": max(dates) if dates else None,
            "program_status": learner.get("program_status") or "Unknown",
            "has_break_in_learning": bool(learner.get("has_break_in_learning")),
            "coach": learner.get("coach") or {"name": None, "email": None},
            # Off-the-job-hours adjustment: the whole-programme summary plus, when
            # a period is selected, that month's specific breakdown + flag.
            "otjh": {
                "adjusted": bool(otjh_summary.get("adjusted")),
                "applied_date": otjh_summary.get("applied_date"),
                "note": otjh_summary.get("note"),
                "flagged_count": otjh_summary.get("flagged_count") or 0,
                "status_counts": otjh_summary.get("status_counts") or {},
                "band_target_h": otjh_summary.get("band_target_h"),
                "band_correct_h": otjh_summary.get("band_correct_h"),
                "flagged_months": otjh_summary.get("flagged_months") or [],
                "month": month_adjustment if isinstance(month_adjustment, dict) else None,
                "month_flagged": bool(month_adjustment.get("flagged")) if isinstance(month_adjustment, dict) else False,
            },
        })

    if search:
        summaries = [item for item in summaries if search in item["name"].lower()]
    if position == "behind":
        summaries = [item for item in summaries if item["gap_hours"] < 0]
    elif position == "ahead":
        summaries = [item for item in summaries if item["gap_hours"] >= 0]

    ordered_periods = sorted(period_values)
    return JsonResponse({
        "learners": summaries,
        "months": [
            {"number": index + 1, "label": month_labels[value]}
            for index, value in enumerate(ordered_periods)
        ],
        "categories": sorted(categories),
        "periods": [
            {"value": value, "label": _period_label(value)}
            for value in ordered_periods
        ],
    })


@require_GET
def learner_profile(request: HttpRequest) -> JsonResponse:
    """Return the rich cross-source profile for any Last_audit learner."""
    learner_key = request.GET.get("learner", "").strip().lower()
    if not learner_key or len(learner_key) > 200:
        return JsonResponse({"error": "A valid learner is required."}, status=400)
    if is_excluded_learner(learner_key, learner_key):
        return JsonResponse({"error": "Learner not found."}, status=404)

    try:
        learner = _load_profile_learner(learner_key)
    except (KeyError, DatabaseError) as error:
        return JsonResponse(
            {"error": "Could not read the learner from Last_audit.", "details": str(error)},
            status=503,
        )

    if learner is None:
        try:
            learner = _fetch_profile_row(learner_key)
        except DatabaseError as error:
            return JsonResponse(
                {"error": "Could not read the learner profile.", "details": str(error)},
                status=503,
            )
    if learner is None:
        return JsonResponse({"error": "Learner not found."}, status=404)

    try:
        sources = _load_profile_sources(
            learner["aptem_id"],
            learner.get("email"),
            learner.get("name"),
        )
    except DatabaseError as error:
        return JsonResponse(
            {"error": "Could not load the learner profile sources.", "details": str(error)},
            status=503,
        )

    training_plan = _training_plan_from_audit(learner.get("training_plan"))

    return JsonResponse({
        "id": _learner_id(learner["name"]),
        "aptem_id": str(learner["aptem_id"]),
        "name": learner["name"],
        "email": learner.get("email"),
        "programme": learner.get("programme_name") or PROGRAMME_NAME,
        "programme_status": (
            sources["programme_status"]
            if sources["programme_status"] not in (None, "", "Unknown")
            else learner.get("programme_status") or "Unknown"
        ),
        "break_in_learning": sources["break_in_learning"],
        "coach": learner.get("coach") or {"name": None, "email": None},
        "planned_hours": sources["learning_delivery"].get("planned_hours"),
        "learning_delivery": sources["learning_delivery"],
        "contracts": sources["contracts"],
        "training_plan": training_plan,
        "skills_radar": sources["skills_radar"],
        "certifications": sources["certifications"],
        "employment": sources["employment"],
        "programme_understanding": sources["programme_understanding"],
    })


def _load_profile_learner(learner_key):
    """Resolve one Aptem-first learner without loading the legacy PCP cohort.

    ``Last_audit.learners`` is the canonical learner list. The Training Plan
    follows the deployed profile and is joined from
    ``Audit.learner_match.aptem_training_plan`` by Aptem ID.
    """
    with connections[resolve(CONN)].cursor() as cursor:
        cursor.execute(
            '''
            select l.aptem_id, l.learner_name, l.learner_email,
                   l.programme_name,
                   coalesce(
                       case
                           when lower(btrim(coalesce(l.programme_status, ''))) in ('', 'unknown') then null
                           else btrim(l.programme_status)
                       end,
                       nullif(btrim(aptem."Program-Status"), '')
                   ) as programme_status,
                   coalesce(nullif(btrim(l.coach_name), ''), nullif(btrim(aptem."OwnerName"), '')) as coach_name,
                   coalesce(nullif(btrim(l.coach_email), ''), nullif(btrim(aptem."OwnerEmail"), '')) as coach_email,
                   lm.aptem_training_plan
            from "Last_audit".learners l
            left join "Audit".learner_match lm
              on lm.aptem_id = l.aptem_id
            left join "LMS"."Aptem_users" aptem
              on aptem."ID" = l.aptem_id
            where l.aptem_id::text = %s
               or lower(l.learner_name) = %s
            order by case when l.aptem_id::text = %s then 0 else 1 end
            limit 1
            ''',
            [learner_key, learner_key, learner_key],
        )
        row = cursor.fetchone()

    if row is None:
        # Some valid Aptem learners (for example a withdrawn learner imported
        # only through the ILR feed) have not yet reached
        # ``Last_audit.learners``.  Keep the Aptem ID as the canonical identity
        # and expose the learner only when an ILR delivery proves that the
        # record belongs to the audit cohort.
        with connections[CONN].cursor() as cursor:
            cursor.execute(
                '''
                select aptem."ID", aptem."FullName", aptem."Email",
                       coalesce(
                           nullif(btrim(aptem."Program Name"), ''),
                           nullif(btrim(aptem."Group"), '')
                       ) as programme_name,
                       nullif(btrim(aptem."Program-Status"), '') as programme_status,
                       nullif(btrim(aptem."OwnerName"), '') as coach_name,
                       nullif(btrim(aptem."OwnerEmail"), '') as coach_email,
                       lm.aptem_training_plan
                from "LMS"."Aptem_users" aptem
                left join "Audit".learner_match lm
                  on lm.aptem_id = aptem."ID"
                where (
                        aptem."ID"::text = %s
                        or lower(btrim(aptem."FullName")) = %s
                        or lower(btrim(aptem."Email")) = %s
                      )
                  and exists (
                        select 1
                        from "Audit".ilr_learning_deliveries ilr
                        where (
                            nullif(btrim(aptem."Email"), '') is not null
                            and lower(btrim(ilr.email)) = lower(btrim(aptem."Email"))
                        ) or (
                            nullif(btrim(aptem."FullName"), '') is not null
                            and lower(btrim(concat_ws(' ', ilr.given_names, ilr.family_name))) =
                                lower(btrim(aptem."FullName"))
                        )
                  )
                order by case when aptem."ID"::text = %s then 0 else 1 end
                limit 1
                ''',
                [learner_key, learner_key, learner_key, learner_key],
            )
            row = cursor.fetchone()

    if row is None:
        return None

    (
        aptem_id, learner_name, learner_email, programme_name,
        programme_status, coach_name, coach_email, training_plan,
    ) = row
    if is_excluded_learner(aptem_id, learner_name):
        return None
    return {
        "aptem_id": aptem_id,
        "name": learner_name,
        "email": learner_email,
        "programme_name": programme_name,
        "programme_status": programme_status,
        "coach": {"name": coach_name, "email": coach_email},
        "training_plan": training_plan,
    }


def _load_profile_sources(aptem_id, learner_email, learner_name=None):
    """Read the profile's external sources without retaining user selection.

    Every query is keyed by this request's learner id/email. The returned data
    is request-local and safe when different learners are viewed concurrently.
    """
    profile_overrides = get_profile_overrides(aptem_id)
    contracts = []
    skill_entries = []
    certifications = []
    employment = None
    cv_employment_candidates = []
    learning_delivery = {}
    programme_understanding = {
        "understanding_programme": None,
        "career_development_progression": None,
    }
    programme_status = "Unknown"
    break_in_learning = {
        "has_break_in_learning": False,
        "last_learning_date": None,
        "expected_return_date": None,
        "has_return_to_learning": False,
        "return_to_learning_date": None,
        "revised_learning_planned_end_date": None,
    }

    with connections[resolve(CONN)].cursor() as cursor:
        ensure_contract_archive_table(cursor)
        cursor.execute(
            '''
            select contracts.id, coalesce(nullif(archive.display_name, ''), contracts.document_name), contracts.status, contracts.date,
                   contracts.learner_signed_date, contracts.fully_signed_date,
                   contracts.requested_date, contracts.program_name,
                   contracts.program_start_date, contracts.planned_end_date,
                   contracts.file, contracts.azure_path,
                   archive.archived_at, archive.archived_by
            from fetching_evidence.aptem_cv_contracts_probe contracts
            left join "Audit".contract_document_archive archive
              on archive.contract_id = contracts.id
            where contracts.learner_id = %s
              and archive.deleted_at is null
            order by contracts.date desc nulls last, contracts.id desc
            ''',
            [aptem_id],
        )
        for row in cursor.fetchall():
            signature_dates = _contract_signature_dates(row[11], row[1])
            document_learner_signed_date = signature_dates.get("learner_signed_date")
            document_fully_signed_date = signature_dates.get("fully_signed_date")
            learner_signed_date = document_learner_signed_date or row[4]
            fully_signed_date = document_fully_signed_date or row[5]
            contracts.append({
                "id": str(row[0]),
                "document_name": row[1] or "Contract",
                "status": row[2] or "Unknown",
                "date": row[3],
                "learner_signed_date": learner_signed_date,
                "fully_signed_date": fully_signed_date,
                "document_learner_signed_date": document_learner_signed_date,
                "document_fully_signed_date": document_fully_signed_date,
                "metadata_learner_signed_date": row[4],
                "metadata_fully_signed_date": row[5],
                "learner_signed_date_source": "document" if document_learner_signed_date else "metadata",
                "fully_signed_date_source": "document" if document_fully_signed_date else "metadata",
                "requested_date": row[6],
                "programme": row[7],
                "programme_start_date": row[8],
                "planned_end_date": row[9],
                "file": f"/audit_api/contracts/{row[0]}/open" if row[11] else row[10],
                "download_file": row[10],
                "azure_path_available": bool(row[11]),
                "archived": bool(row[12]),
                "archived_at": row[12],
                "archived_by": row[13],
            })

        cursor.execute(
            '''
            select program_status, "Break in learning"
            from fetching_evidence.aptem_cv_contracts_probe
            where learner_id = %s and source <> 'audit_upload'
            order by fetched_at desc nulls last, id desc
            limit 1
            ''',
            [aptem_id],
        )
        status_row = cursor.fetchone()
        if status_row:
            programme_status = status_row[0] or "Unknown"
            break_value = status_row[1]
            if isinstance(break_value, str):
                try:
                    break_value = json.loads(break_value)
                except ValueError:
                    break_value = None
            if isinstance(break_value, dict):
                break_in_learning = {
                    "has_break_in_learning": bool(break_value.get("has_break_in_learning")),
                    "last_learning_date": break_value.get("last_learning_date"),
                    "expected_return_date": break_value.get("expected_return_date"),
                    "has_return_to_learning": bool(break_value.get("has_return_to_learning")),
                    "return_to_learning_date": break_value.get("return_to_learning_date"),
                    "revised_learning_planned_end_date": break_value.get("revised_learning_planned_end_date"),
                }
            if str(programme_status).strip().lower() == "onbreak":
                break_in_learning["has_break_in_learning"] = True
        break_in_learning = apply_break_overrides(break_in_learning, profile_overrides)

        cursor.execute(
            '''
            select
                coalesce(
                    "Programme understanding" ->> 'understanding_programme',
                    "Programme understanding" -> 'raw' ->> 'ExtendedILRModel_UnderstandingProgramme'
                ),
                coalesce(
                    "Programme understanding" ->> 'career_development_progression',
                    "Programme understanding" -> 'raw' ->> 'ExtendedILRModel_CareerDevelopmentProgression'
                )
            from fetching_evidence.aptem_cv_contracts_probe
            where learner_id = %s
              and "Programme understanding" is not null
            order by fetched_at desc nulls last, id desc
            limit 1
            ''',
            [aptem_id],
        )
        understanding_row = cursor.fetchone()
        if understanding_row:
            programme_understanding = {
                "understanding_programme": understanding_row[0] or None,
                "career_development_progression": understanding_row[1] or None,
            }

        cursor.execute(
            '''
            select distinct on (lower(btrim(characteristic_name)))
                   characteristic_name, assessed_level,
                   raw -> 'score' ->> 'achieved',
                   raw -> 'score' ->> 'maximum'
            from fetching_evidence.aptem_skills_radar_probe
            where learner_id = %s
              and coalesce(raw -> 'score' ->> 'achieved', assessed_level::text) is not null
            order by lower(btrim(characteristic_name)), fetched_at desc, id desc
            ''',
            [aptem_id],
        )
        for characteristic, assessed_level, achieved_value, maximum_value in cursor.fetchall():
            # Aptem can return programme Duties in the same payload as the
            # actual K/S/B characteristics. Duties have no KSB code and were
            # previously inflating Skills (for example 11 skills became 33)
            # and filling the radar with paragraph-length axis labels.
            if re.match(r"^\s*duty\s+\d+\b", str(characteristic or ""), re.IGNORECASE):
                continue
            domain, field = _skill_radar_characteristic(characteristic)
            score, maximum = _skill_radar_score_values(
                assessed_level,
                achieved_value,
                maximum_value,
            )
            characteristic_text = str(characteristic or "").strip()
            skill_entries.append({
                "skill": characteristic_text or domain,
                "domain": domain,
                "ksb_codes": _skill_radar_codes(characteristic_text),
                "knowledge": score if field == "knowledge" else None,
                "skill_score": score if field == "skill_score" else None,
                "behaviour": score if field == "behaviour" else None,
                "maximum": maximum,
            })

        if not skill_entries and str(programme_status).strip().lower() == "withdrawn":
            cursor.execute(
                '''
                select category, code, title, level
                from fetching_evidence.learner_ksbs
                where learner_id = %s and level is not null
                order by category, code
                ''',
                [aptem_id],
            )
            skill_entries.extend(_skill_radar_snapshot_entries(cursor.fetchall()))

        cursor.execute(
            '''
            select certifications, employment_details
            from fetching_evidence.aptem_cv_certifications
            where learner_id = %s
            order by updated_at desc nulls last, id desc
            ''',
            [aptem_id],
        )
        seen_certifications = set()
        for certification_value, employment_value in cursor.fetchall():
            if isinstance(certification_value, str):
                try:
                    certification_value = json.loads(certification_value)
                except ValueError:
                    certification_value = []
            if isinstance(certification_value, list):
                for certification in certification_value:
                    if not isinstance(certification, dict):
                        continue
                    key = (
                        str(certification.get("name") or "").strip().lower(),
                        str(certification.get("issuer") or "").strip().lower(),
                    )
                    if key in seen_certifications or not key[0]:
                        continue
                    seen_certifications.add(key)
                    certifications.append(certification)
            cv_employment_terms = _cv_employment_terms(employment_value)
            if cv_employment_terms:
                cv_employment_candidates.append(cv_employment_terms)

        cursor.execute(
            '''
            select learner_employer_details
            from fetching_evidence.aptem_cv_contracts_probe
            where learner_id = %s
              and learner_employer_details is not null
              and learner_employer_details <> '{}'::jsonb
            order by fetched_at desc nulls last, id desc
            limit 1
            ''',
            [aptem_id],
        )
        employer_row = cursor.fetchone()
        if employer_row:
            employment = _employer_details_from_contract_profile(employer_row[0])
        employment = _merge_matching_cv_employment_terms(employment, cv_employment_candidates)

        cursor.execute(
            '''
            select "Levy or Not"
            from "LMS"."Aptem_users"
            where "ID" = %s
            limit 1
            ''',
            [aptem_id],
        )
        levy_row = cursor.fetchone()
        levy_status = _normalise_levy_status(levy_row[0] if levy_row else None)
        if levy_status:
            employment = dict(employment or {})
            employment["levy_status"] = levy_status

        if learner_email or learner_name:
            learner_email = str(learner_email or "").strip()
            learner_name = str(learner_name or "").strip()
            cursor.execute(
                '''
                select learn_ref_number, planned_hours, otj_actual_hours,
                       min(learn_start_date) over () as original_programme_start_date,
                       learn_plan_end_date, completion_status,
                       max(learn_actual_end_date) over () as actual_end_date,
                       nullif(
                           concat_ws(', ',
                               nullif(btrim(address_line_1), ''),
                               nullif(btrim(address_line_2), ''),
                               nullif(btrim(address_line_3), ''),
                               nullif(btrim(address_line_4), '')
                           ),
                           ''
                       ) as learner_address,
                       nullif(btrim(postcode), '') as learner_postcode,
                       nullif(btrim(delivery_location_postcode), '') as employer_postcode
                from "Audit".ilr_learning_deliveries
                where planned_hours is not null
                  and (
                    (%s <> '' and lower(btrim(email)) = lower(%s))
                    or
                    (%s <> '' and lower(btrim(concat_ws(' ', given_names, family_name))) = lower(%s))
                  )
                order by
                    case when %s <> '' and lower(btrim(email)) = lower(%s) then 0 else 1 end,
                    aim_seq_number, updated_at desc nulls last, id desc
                limit 1
                ''',
                [
                    learner_email, learner_email,
                    learner_name, learner_name,
                    learner_email, learner_email,
                ],
            )
            delivery = cursor.fetchone()
            if delivery:
                learning_delivery = {
                    "learner_reference": delivery[0],
                    "planned_hours": delivery[1],
                    "actual_hours": delivery[2],
                    "start_date": delivery[3],
                    "planned_end_date": delivery[4],
                    "completion_status": delivery[5],
                    "actual_end_date": delivery[6],
                    "last_learning_evidence_date": break_in_learning.get("last_learning_date") or delivery[6],
                    "learner_address": delivery[7],
                    "learner_postcode": delivery[8],
                    "employer_postcode": delivery[9],
                    "first_evidence_date": None,
                    "first_evidence_items": [],
                    "archived_evidence_items": [],
                    "last_learning_evidence_items": [],
                    "break_evidence_items": [],
                }
                ensure_evidence_override_table(cursor)
                cursor.execute(
                    '''
                    with raw_candidates as (
                        select
                            item ->> 'id' as evidence_id,
                            item ->> 'name' as evidence_name,
                            item ->> 'component_name' as component_name,
                            item ->> 'kind' as evidence_kind,
                            item ->> 'status' as evidence_status,
                            item ->> 'file' as evidence_file,
                            item ->> 'content' as evidence_content,
                            substring(item ->> 'created_date' from 1 for 10)::date as evidence_date
                        from fetching_evidence.learner_evidence learner_evidence
                        cross join lateral jsonb_array_elements(
                            case
                                when jsonb_typeof(learner_evidence.evidence) = 'array'
                                    then learner_evidence.evidence
                                else '[]'::jsonb
                            end
                        ) item
                        where learner_evidence.learner_id = %s
                          and ltrim(lower(coalesce(item ->> 'name', ''))) not like 'welcome%%'
                          and ltrim(lower(coalesce(item ->> 'component_name', ''))) not like 'welcome%%'
                          and coalesce(item ->> 'created_date', '') ~ '^\\d{4}-\\d{2}-\\d{2}'
                          and substring(item ->> 'created_date' from 1 for 10)::date >= %s
                    ), candidates as (
                        select distinct on (evidence_id)
                            evidence_id, evidence_name, component_name, evidence_kind,
                            evidence_status, evidence_file, evidence_content, evidence_date
                        from raw_candidates
                        order by evidence_id, evidence_date
                    )
                    select candidates.evidence_id, candidates.evidence_name,
                           candidates.component_name, candidates.evidence_kind,
                           candidates.evidence_status, candidates.evidence_file,
                           candidates.evidence_content,
                           coalesce(overrides.evidence_date, candidates.evidence_date) as evidence_date,
                           overrides.archived_at is not null as archived,
                           overrides.deleted_at is not null as deleted,
                           false as uploaded,
                           null::bigint as source_activity_id,
                           null::text as source_activity_month,
                           null::text as source_activity_category
                    from candidates
                    left join "Audit".learner_evidence_overrides overrides
                      on overrides.learner_id = %s
                     and overrides.is_uploaded = false
                     and overrides.source_evidence_id::text = candidates.evidence_id
                    union all
                    select uploads.evidence_id, uploads.document_name,
                           uploads.component_name, uploads.evidence_kind,
                           uploads.evidence_status, null, null, uploads.evidence_date,
                           uploads.archived_at is not null as archived,
                           uploads.deleted_at is not null as deleted,
                           true as uploaded,
                           uploads.source_activity_id,
                           uploads.source_activity_month,
                           uploads.source_activity_category
                    from "Audit".learner_evidence_overrides uploads
                    where uploads.learner_id = %s and uploads.is_uploaded = true
                    order by evidence_date, evidence_id
                    ''',
                    [aptem_id, delivery[3], aptem_id, aptem_id],
                )
                evidence_rows = cursor.fetchall()
                if evidence_rows:
                    evidence_items = [
                        {
                            "id": row[0],
                            "name": row[1] or "Untitled evidence",
                            "component_name": row[2] or "",
                            "kind": row[3] or "",
                            "status": row[4] or "",
                            "file": f"/audit_api/evidence/{quote(str(row[0]), safe='')}/open?learner_id={aptem_id}" if row[0] and row[11] is None else None,
                            "content": row[6],
                            "date": row[7],
                            "archived": bool(row[8]),
                            "deleted": bool(row[9]),
                            "uploaded": bool(row[10]),
                            "source_activity_id": row[11],
                            "source_activity_month": row[12],
                            "source_activity_category": row[13],
                        }
                        for row in evidence_rows
                    ]
                    first_date, first_items, archived_items = _partition_evidence_items(evidence_items)
                    learning_delivery["archived_evidence_items"] = archived_items
                    learning_delivery["first_evidence_date"] = first_date
                    learning_delivery["first_evidence_items"] = first_items
                    learning_delivery["last_learning_evidence_items"] = _last_learning_evidence_items(
                        evidence_items,
                        learning_delivery.get("last_learning_evidence_date"),
                    )
                    learning_delivery["break_evidence_items"] = _break_evidence_items(
                        evidence_items,
                        break_in_learning.get("return_to_learning_date"),
                    )

    # Keep every assessed characteristic as its own radar axis. A programme can
    # legitimately have two Knowledge rows in the same domain (for example
    # Strategic Project Management K1 and K30); grouping by domain + dimension
    # used to overwrite one of those rows.
    skills_radar = sorted(skill_entries, key=_skill_radar_entry_sort_key)
    employment, learning_delivery = apply_profile_overrides(
        employment,
        learning_delivery,
        profile_overrides,
    )
    return {
        "contracts": contracts,
        "skills_radar": skills_radar,
        "certifications": certifications,
        "employment": employment,
        "learning_delivery": learning_delivery,
        "programme_understanding": programme_understanding,
        "programme_status": programme_status,
        "break_in_learning": break_in_learning,
    }


def _partition_evidence_items(evidence_items):
    """Keep Aptem's original first evidence fixed; uploads may replace it.

    Archiving the original source evidence must not promote Aptem's second
    evidence. An auditor-uploaded item is the only replacement candidate.
    """
    archived_items = [
        item for item in evidence_items if item["archived"] and not item["deleted"]
    ]
    source_items = [item for item in evidence_items if not item["uploaded"]]
    original_source_date = min((item["date"] for item in source_items), default=None)
    qualifying_items = [
        item for item in evidence_items
        if not item["archived"]
        and not item["deleted"]
        and (
            (
                item["uploaded"]
                and item.get("component_name") not in {
                    "Break in learning evidence",
                    "Last date of learning evidence",
                    "Return to learning evidence",
                }
            )
            or (original_source_date is not None and item["date"] == original_source_date)
        )
    ]
    first_date = min((item["date"] for item in qualifying_items), default=None)
    first_items = [item for item in qualifying_items if item["date"] == first_date]
    return first_date, first_items, archived_items


def _break_evidence_items(evidence_items, evidence_date):
    """Return active evidence recorded on one significant break date."""
    target_date = _iso_date(evidence_date)
    if target_date is None:
        return []
    return [
        item for item in evidence_items
        if not item["archived"]
        and not item["deleted"]
        and _iso_date(item["date"]) == target_date
    ]


def _last_learning_evidence_items(evidence_items, last_learning_date):
    """Use same-day evidence, falling back to the latest earlier learning day."""
    target_date = _iso_date(last_learning_date)
    if target_date is None:
        return []
    dated_items = [
        item for item in evidence_items
        if _iso_date(item["date"]) is not None
        and _iso_date(item["date"]) <= target_date
    ]
    matching_date = max(
        (_iso_date(item["date"]) for item in dated_items),
        default=None,
    )
    if matching_date is None:
        return []
    return [
        item for item in dated_items
        if _iso_date(item["date"]) == matching_date
        and not item["archived"]
        and not item["deleted"]
    ]


def _employer_details_from_contract_profile(value):
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except ValueError:
            return None
    if not isinstance(value, dict):
        return None

    employer = value.get("employer")
    if not isinstance(employer, dict):
        employer = {}
    raw = value.get("raw")
    if not isinstance(raw, dict):
        raw = {}
    manager = employer.get("manager")
    if not isinstance(manager, dict):
        manager = {}
    address = employer.get("address")
    if not isinstance(address, dict):
        address = {}

    employer_name = employer.get("name") or raw.get("UserEmployer_Organization")
    job_title = employer.get("job_title") or raw.get("UserILRSummary_JobTitle")
    workplace_address = address.get("formatted") or raw.get("UserEmployer_Address")
    manager_name = manager.get("name") or raw.get("UserEmployer_ManagerName")
    manager_email = manager.get("email") or raw.get("UserEmployer_ManagerEmail")
    manager_phone = manager.get("phone_number") or raw.get("UserEmployer_ManagerPhone")
    if not any((employer_name, job_title, workplace_address, manager_name, manager_email, manager_phone)):
        return None
    return {
        "employer_name": employer_name or None,
        "job_title": job_title or None,
        "workplace_address": workplace_address or None,
        "employment_start_date": None,
        "contracted_hours_per_week": None,
        "employment_type": None,
        "working_pattern": None,
        "line_manager": {
            "name": manager_name or None,
            "email": manager_email or None,
            "phone": manager_phone or None,
            "job_title": None,
        },
    }


def _cv_employment_terms(value):
    """Read the two contract terms that only exist in the extracted CV data."""
    if isinstance(value, str):
        try:
            decoded = json.loads(value)
        except ValueError:
            return None
        if decoded == value:
            return None
        return _cv_employment_terms(decoded)

    if isinstance(value, list):
        for item in value:
            terms = _cv_employment_terms(item)
            if terms:
                return terms
        return None

    if not isinstance(value, dict):
        return None

    if value.get("section_found", True):
        employer_name = _cv_optional_text(value.get("employer_name"))
        start_date = _cv_optional_text(value.get("employment_start_date"))
        contracted_hours = _to_float(value.get("contracted_hours_per_week"))
        if start_date is not None or contracted_hours is not None:
            return {
                "employer_name": employer_name,
                "employment_start_date": start_date,
                "contracted_hours_per_week": contracted_hours,
            }

    for nested in value.values():
        terms = _cv_employment_terms(nested)
        if terms:
            return terms
    return None


def _cv_optional_text(value):
    if value is None:
        return None
    text = str(value).strip()
    if text.lower() in {"", "none", "null", "n/a", "-"}:
        return None
    return text


def _skill_radar_characteristic(value):
    text = str(value or "").strip() or "Skill"
    labelled_match = re.match(
        r"Understanding of (.+?) \((Knowledge|Skill|Behaviour)\)",
        text,
        re.IGNORECASE,
    )
    if labelled_match:
        score_type = labelled_match.group(2).lower()
        field = {
            "knowledge": "knowledge",
            "skill": "skill_score",
            "behaviour": "behaviour",
        }[score_type]
        return labelled_match.group(1).strip(), field

    code_match = re.search(r"(?:^|[\s(:.\-])([KSB])\d+\b", text, re.IGNORECASE)
    prefix = code_match.group(1).upper() if code_match else "S"
    field = {
        "K": "knowledge",
        "S": "skill_score",
        "B": "behaviour",
    }[prefix]
    # A number of Aptem standards use ``K01 Description`` rather than
    # ``K01: Description``. Treat both forms identically so the full KSB text
    # never becomes the chart's category label.
    direct_code = re.match(
        r"\s*[KSB]0*\d+\s*(?::|[.\-])?\s+(.+)",
        text,
        re.IGNORECASE,
    )
    return (_skill_radar_text_category(direct_code.group(1)) if direct_code else text), field


_SKILL_RADAR_CATEGORY_RULES = (
    (r"works flexibly|adapts? to circumstances", "Adaptability"),
    (r"works collaboratively|builds strong relationships", "Collaboration"),
    (r"accountability and ownership|ownership of (?:their|the) tasks", "Accountability"),
    (r"operates professionally|integrity and confidentiality", "Professionalism"),
    (r"learning opportunities|continuous professional development", "Continuous development"),
    (r"differences between projects and business as usual|alignment between the project and organisational objectives|interdependencies between project, programme|project context|project governance structure|functional, matrix and project structures|roles and responsibilities within a project|life cycle approaches|business case|project management plan", "Project context & governance"),
    (r"define,? record,? integrate,? deliver,? and manage scope|configuration management and change control|change control processes?|management of project scope", "Scope & change"),
    (r"stakeholders?|communication techniques|managing conflict|working collaboratively|influence and negotiate|resolve conflict|adapt communications?|project vision", "Stakeholders & communication"),
    (r"information management|technology and software|digital tools?|presentation tools", "Information & technology"),
    (r"estimating methods?|earned value|project scheduling|schedule activities|integrated schedules|allocation and management of resources|project budgets?|resource management|manages resources|resources through the project|critical path|approved project budget", "Planning, cost & resources"),
    (r"project risk and issue|risk management plan|project risks? and issues?|mitigate risks?", "Risk & issues"),
    (r"procurement strategies|quality requirements|quality management plan|quality control", "Procurement & quality"),
    (r"evaluating project success|lessons learned|continual improvement", "Evaluation & improvement"),
    (r"regulations? and legislation|relevant legislation|sustainability|net carbon|ethical and inclusive|codes? of practice|ethical guidance", "Compliance & sustainability"),
    (r"monitoring and reporting|track,? interpret and report|collate and analyse information|use data to inform|underpinning data", "Monitoring & reporting"),
    (r"fundamentals of marketing theory|marketing process", "Marketing fundamentals"),
    (r"brand positioning|corporate reputation", "Brand management"),
    (r"customer relationship management|stakeholder management.*customer", "Stakeholder & CRM"),
    (r"business and sector|vision and value", "Business & sector"),
    (r"wider business objectives", "Business objectives"),
    (r"target audience.*decision", "Customer behaviour"),
    (r"legal.*regulatory|compliance frameworks?|data protection", "Legal & compliance"),
    (r"principles of effective market research", "Market research"),
    (r"product development|product/service portfolios?", "Product development"),
    (r"routes? to market|marketing landscape", "Routes to market"),
    (r"communications?\s+channels? and media", "Communication channels"),
    (r"coordinate and maintain.*marketing channels", "Marketing channels"),
    (r"tactical campaigns?|smart objectives?", "Campaign planning"),
    (r"production and distribution.*marketing materials?", "Marketing materials"),
    (r"creative and effective communications?|write and proofread", "Marketing communications"),
    (r"engage and collaborate.*(?:clients?|stakeholders?)|across departments", "Stakeholder collaboration"),
    (r"project and time management", "Project & time management"),
    (r"coordinate several marketing campaigns", "Campaign coordination"),
    (r"liaise with.*stakeholders?|manage.*stakeholders?.*suppliers?", "Stakeholder management"),
    (r"project budgets?|budget", "Budget management"),
    (r"assimilate and analyse data|data and information.*range of sources", "Data analysis"),
    (r"effectiveness of marketing campaigns", "Campaign evaluation"),
    (r"data and research.*derive insights|insights.*future campaigns", "Research insights"),
    (r"business systems? and software", "Business systems"),
    (r"appropriate technologies|web analytics|social media.*crm", "Marketing technology"),
    (r"tenacious and driven|projects? through to completion", "Drive & resilience"),
    (r"self.?starter|adaptable approach|changing work priorities", "Initiative & adaptability"),
    (r"creative and analytical mind|new ways of doing", "Creative thinking"),
    (r"ideas and solutions", "Problem solving"),
    (r"learn from mistakes|improve.*performance", "Continuous improvement"),
    (r"professionalism|reliability and dependability", "Professionalism"),
    (r"collaborative approach|showing empathy", "Collaboration & empathy"),
    (r"ethical behaviour|equality.*diversity", "Ethics & inclusion"),
)


def _skill_radar_text_category(value):
    """Build a concise chart category when Aptem only supplies full KSB text."""
    text = " ".join(str(value or "").split()).strip()
    lowered = text.lower()
    for pattern, category in _SKILL_RADAR_CATEGORY_RULES:
        if re.search(pattern, lowered):
            return category

    concise = re.sub(
        r"^(?:i\s+)?(?:can|understand|know|am able to|have|work with|demonstrate|come up with)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    words = concise.rstrip(" .").split()
    if len(words) > 5:
        concise = " ".join(words[:5])
    return concise[:60].strip().capitalize() or "Competency"


def _skill_radar_codes(value):
    """Return each KSB code once, in the order Aptem supplied it."""
    codes = []
    seen = set()
    for match in re.finditer(r"\b([KSB]\d+)\b", str(value or ""), re.IGNORECASE):
        code = match.group(1).upper()
        if code not in seen:
            seen.add(code)
            codes.append(code)
    return codes


def _skill_radar_score_values(assessed_level, achieved_value, maximum_value):
    """Return the Aptem score using the maximum supplied for this characteristic."""
    try:
        maximum = float(maximum_value)
    except (TypeError, ValueError):
        maximum = 8.0
    if maximum <= 0:
        maximum = 8.0

    try:
        achieved = float(achieved_value)
    except (TypeError, ValueError):
        try:
            achieved = float(assessed_level)
        except (TypeError, ValueError):
            achieved = 0.0

    achieved = max(0.0, min(maximum, achieved))
    return (
        int(achieved) if achieved.is_integer() else achieved,
        int(maximum) if maximum.is_integer() else maximum,
    )


def _skill_radar_snapshot_entries(rows):
    """Normalise the retained KSB snapshot used when a withdrawn Aptem probe is absent."""
    entries = []
    for category, code, title, level in rows:
        score = _skill_radar_level_score(level)
        if score is None:
            continue
        code = str(code or "").strip().upper()
        category_text = str(category or "").strip().lower()
        prefix = code[:1]
        if category_text.startswith("know") or prefix == "K":
            field = "knowledge"
        elif category_text.startswith("behav") or prefix == "B":
            field = "behaviour"
        else:
            field = "skill_score"
        title = str(title or "").strip()
        entries.append({
            "skill": f"{code}: {title}" if code and title else title or code,
            "domain": _skill_radar_text_category(title),
            "ksb_codes": [code] if re.fullmatch(r"[KSB]0*\d+", code) else [],
            "knowledge": score if field == "knowledge" else None,
            "skill_score": score if field == "skill_score" else None,
            "behaviour": score if field == "behaviour" else None,
            "maximum": 8,
        })
    return entries


def _skill_radar_level_score(value):
    text = str(value or "").strip().lower()
    for label, score in (
        ("mastery", 8), ("expert", 7), ("proficient", 6),
        ("consistently", 5), ("frequently", 4), ("occasionally", 3),
        ("rarely", 2), ("never", 1),
    ):
        if text.startswith(label):
            return score
    return None


def _skill_radar_entry_sort_key(entry):
    fields = {
        "knowledge": 0 if entry.get("knowledge") is not None else None,
        "skill_score": 1 if entry.get("skill_score") is not None else None,
        "behaviour": 2 if entry.get("behaviour") is not None else None,
    }
    dimension = next((value for value in fields.values() if value is not None), 3)
    first_code = (entry.get("ksb_codes") or [""])[0]
    code_match = re.fullmatch(r"[KSB](\d+)", first_code)
    return (
        dimension,
        int(code_match.group(1)) if code_match else 10**9,
        str(entry.get("domain") or "").lower(),
        str(entry.get("skill") or "").lower(),
    )


def _normalise_levy_status(value):
    text = re.sub(r"[\s_-]+", "", str(value or "")).lower()
    if text == "levy":
        return "Levy"
    if text in {"nonlevy", "notlevy"}:
        return "Non-Levy"
    return None


def _skill_radar_sort_key(value):
    text = str(value or "")
    code_match = re.search(r"(?:^|[\s(:.\-])([KSB])(\d+)\b", text, re.IGNORECASE)
    if not code_match:
        return (3, 0, text.lower())
    prefix_order = {"K": 0, "S": 1, "B": 2}
    return (prefix_order[code_match.group(1).upper()], int(code_match.group(2)), text.lower())


def _normalise_employer_name(value):
    text = _cv_optional_text(value)
    if text is None:
        return None
    words = re.findall(r"[a-z0-9]+", text.lower().replace("&", " and "))
    suffixes = {"limited", "ltd"}
    while words and words[-1] in suffixes:
        words.pop()
    return " ".join(words) or None


def _merge_matching_cv_employment_terms(employment, candidates):
    if not isinstance(employment, dict):
        return employment
    employer_name = employment.get("employer_name")
    matching_terms = next(
        (
            candidate
            for candidate in candidates
            if _normalise_employer_name(candidate.get("employer_name"))
            == _normalise_employer_name(employer_name)
            and _normalise_employer_name(employer_name) is not None
        ),
        None,
    )
    if matching_terms is None:
        return employment
    if employment.get("employment_start_date") is None:
        employment["employment_start_date"] = matching_terms.get("employment_start_date")
    if employment.get("contracted_hours_per_week") is None:
        employment["contracted_hours_per_week"] = matching_terms.get("contracted_hours_per_week")
    return employment


def _period_label(value):
    try:
        return datetime.datetime.strptime(value, "%Y-%m").strftime("%B %Y")
    except ValueError:
        return value


@require_GET
def learner_activities(request: HttpRequest) -> JsonResponse:
    """Paginated, filtered activity rows. Matches ``getLearnerActivities()``."""
    try:
        limit = _integer_param(request, "limit", 25, 1, 200)
        offset = _integer_param(request, "offset", 0, 0, 1_000_000)
        search = request.GET.get("search", "").strip().lower()[:100]
        learner_filter = request.GET.get("learner", "").strip().lower()
        learner_search = request.GET.get("learner_search", "").strip().lower()[:100]
        category = request.GET.get("category", "").strip().lower()
        period = request.GET.get("period", "").strip()
        if period and (len(period) != 7 or period[4] != "-" or not period.replace("-", "").isdigit()):
            raise ValueError("period must use YYYY-MM format")
    except (TypeError, ValueError) as error:
        return JsonResponse({"error": "Invalid query parameters", "details": str(error)}, status=400)

    try:
        learners_raw = _load_rows()
    except (KeyError, DatabaseError) as error:
        return JsonResponse(
            {"error": "Could not read Audit.learner_match from the enrolment database.", "details": str(error)},
            status=503,
        )

    in_scope = []
    for learner in learners_raw:
        learner_id = _learner_id(learner["name"])
        if learner_filter and learner_filter not in (learner_id, str(learner["aptem_id"])):
            continue
        if learner_search and learner_search not in learner["name"].lower():
            continue
        in_scope.append(learner)

    rows = [item for learner in in_scope for item in learner["activities"]]
    if period:
        rows = [item for item in rows if item["activity_period"] == period]
    if category:
        rows = [item for item in rows if (item["activity_category"] or "").lower() == category]
    if search:
        rows = [
            item for item in rows
            if search in (item["activity_unit"] or "").lower()
            or search in (item["plan_id"] or "").lower()
            or search in (item["activity_description"] or "").lower()
            or search in (item["learner"] or "").lower()
        ]

    rows.sort(key=lambda item: (item["activity_date"] is None, item["activity_date"] or "", item["learner"]))

    # Header totals use the curated month.hours (authoritative), not the noisy
    # per-activity tracked time — summed over the in-scope learners / period.
    planned_total = 0.0
    actual_total = 0.0
    for learner in in_scope:
        if period:
            entry = learner["month_hours"].get(period)
            if entry:
                planned_total += entry["planned"]
                actual_total += entry["completed"]
        else:
            planned_total += sum(value["planned"] for value in learner["month_hours"].values())
            actual_total += sum(value["completed"] for value in learner["month_hours"].values())
    planned_total = round(planned_total, 2)
    actual_total = round(actual_total, 2)
    total = len(rows)
    # For the single-learner + single-month view (the monthly report/journal),
    # surface that month's OTJH adjustment (hour breakdown + flag).
    otjh_month = None
    if period and len(in_scope) == 1:
        candidate = in_scope[0].get("otjh_months", {}).get(period)
        if isinstance(candidate, dict):
            # Attach the programme-level applied timestamp so the monthly report
            # can show when the OTJH adjustment was run.
            summary = in_scope[0].get("otjh_summary") or {}
            otjh_month = {**candidate, "applied_date": summary.get("applied_date"), "note": summary.get("note")}
    return JsonResponse({
        "items": rows[offset : offset + limit],
        "total": total,
        "planned_total": planned_total,
        "actual_total": actual_total,
        "limit": limit,
        "offset": offset,
        "otjh": otjh_month,
    })


@require_GET
def activity_learners(request: HttpRequest) -> JsonResponse:
    """All learners who have a given activity (by component id), each with their
    own log entry on it. Powers the "Learner activity records" table on the
    activity-detail page — a per-activity, cross-learner view. Same response
    shape as ``learner_activities`` so the UI table can consume it unchanged."""
    component = request.GET.get("component", "").strip()
    search = request.GET.get("search", "").strip().lower()[:100]
    if not component:
        return JsonResponse({"items": [], "total": 0, "planned_total": 0.0, "actual_total": 0.0, "limit": 0, "offset": 0})

    try:
        learners_raw = _load_rows()
    except (KeyError, DatabaseError) as error:
        return JsonResponse(
            {"error": "Could not read Audit.learner_match from the enrolment database.", "details": str(error)},
            status=503,
        )

    rows = []
    for learner in learners_raw:
        if search and search not in learner["name"].lower():
            continue
        for item in learner["activities"]:
            if str(item["plan_id"]) == component:
                rows.append(item)

    rows.sort(key=lambda item: (item["actual_lms_hours"] is None, -(item["actual_lms_hours"] or 0), item["learner"]))
    planned_total = round(sum(item["planned_hours"] or 0 for item in rows), 2)
    actual_total = round(sum(item["actual_lms_hours"] or 0 for item in rows), 2)
    return JsonResponse({
        "items": rows,
        "total": len(rows),
        "planned_total": planned_total,
        "actual_total": actual_total,
        "limit": len(rows),
        "offset": 0,
    })


@require_GET
def attendance_session(request: HttpRequest) -> JsonResponse:
    """Everything about one live session, given any one learner's attendance key.

    Returns (a) every learner who attended the *same* session — matched by the
    key's date + normalised group, not the per-learner key — as activity rows
    the "Learner activity records" table can render, and (b) the session's
    recordings (LMS lessons whose title-date equals the session date and whose
    course name overlaps the group), for the content-preview panel."""
    key = request.GET.get("key", "").strip()
    parsed = _parse_attendance_key(key)
    if not parsed:
        return JsonResponse({"error": "key must look like <learner_id>_<YYYY-MM-DD>_<group>"}, status=400)
    _learner_part, date, group_raw = parsed
    norm_group = _norm_group(group_raw)

    try:
        learners_raw = _load_rows()
    except (KeyError, DatabaseError) as error:
        return JsonResponse(
            {"error": "Could not read Audit.learner_match from the enrolment database.", "details": str(error)},
            status=503,
        )

    # (a) attendees — same date + same normalised group, one row per learner.
    attendees = []
    for learner in learners_raw:
        for item in learner["activities"]:
            if item.get("session_date") == date and item.get("session_group") == norm_group:
                attendees.append(item)
    attendees.sort(key=lambda item: item["learner"])

    # (b) recordings — LMS lessons whose title is dated to this session day. On a
    # given day more than one course can run (e.g. a PPC session and a PMO
    # session), so rather than a fixed word-overlap threshold we assign each
    # dated recording to whichever of that day's sessions its course name best
    # matches, and keep the ones that land on THIS session. This copes with
    # topical titles that share only one word with the group ("… PMO …").
    sessions_on_date = {}  # normalised group -> significant-token set
    for learner in learners_raw:
        for item in learner["activities"]:
            group = item.get("session_group")
            if item.get("session_date") == date and group and group not in sessions_on_date:
                sessions_on_date[group] = _significant_tokens(group)

    recordings = {}
    for learner in learners_raw:
        for item in learner["activities"]:
            source_url = item.get("source_url")
            if not source_url:
                continue
            title = item.get("activity_unit")
            if _title_session_date(title) != date:
                continue
            course_tokens = _significant_tokens(_recording_course(title))
            scores = {group: len(course_tokens & tokens) for group, tokens in sessions_on_date.items()}
            best = max(scores.values(), default=0)
            # Keep only if this session is (one of) the best match and it is a
            # real match (>0 shared words), so unrelated same-day courses split.
            if best <= 0 or scores.get(norm_group, 0) != best:
                continue
            component = item.get("plan_id")
            if component in recordings:
                continue
            recordings[component] = {
                "component_id": component,
                "title": _clean_title(title),
                "preview_url": source_url,
                "week": item.get("week"),
            }
    recordings_list = sorted(recordings.values(), key=lambda rec: rec["title"])

    planned_total = round(sum(item["planned_hours"] or 0 for item in attendees), 2)
    actual_total = round(sum(item["actual_lms_hours"] or 0 for item in attendees), 2)
    return JsonResponse({
        "session": {
            "date": date,
            "group": group_raw,
            "group_label": norm_group,
            "module": attendees[0]["activity_unit"] if attendees else None,
        },
        "recordings": recordings_list,
        "items": attendees,
        "total": len(attendees),
        "planned_total": planned_total,
        "actual_total": actual_total,
        "limit": len(attendees),
        "offset": 0,
    })


@require_GET
def quiz_attempt(request: HttpRequest) -> JsonResponse:
    """Return one learner's graded quiz body for a quiz component, read from the
    Audit.learner_match.quiz_attempts jsonb column (keyed by component id,
    attempted quizzes only). Shape: {title, status, quiz_body: {questions: [...]}}.
    Uses the jsonb ``->`` operator to extract just the one quiz rather than
    hauling the whole (large) column back."""
    learner = request.GET.get("learner", "").strip().lower()
    component = request.GET.get("component", "").strip()
    if not learner or not component:
        return JsonResponse({"error": "learner and component are required"}, status=400)
    try:
        with connections[resolve(CONN)].cursor() as cur:
            cur.execute(
                '''
                select quiz_attempts -> %s
                from "Audit".learner_match
                where lower(learner_name) = %s
                  and programme_structure ->> 'programme' = %s
                limit 1
                ''',
                [component, learner, PROGRAMME_NAME],
            )
            row = cur.fetchone()
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not read quiz attempts.", "details": str(error)}, status=503)

    attempt = row[0] if row else None
    if isinstance(attempt, str):
        try:
            attempt = json.loads(attempt)
        except ValueError:
            attempt = None
    return JsonResponse({"component_id": component, "attempt": attempt or None})


# --- manual auditor annotations (KSBs + planned hours) per activity ---------
#
# programme_structure carries no KSB mapping and no auditor-set planned hours,
# so auditors enter those by hand. They are stored once per activity (keyed by
# component id) in a small table this module owns. The table is created on first
# use with `create table if not exists`, mirroring how audit_api/views.py
# provisions its `monthly_audit_signoffs` table.

def _ensure_annotation_table(cur):
    cur.execute(
        '''
        create table if not exists "Audit".activity_annotations (
            component_id text primary key,
            planned_hours numeric,
            mapped_ksbs text,
            updated_by text,
            updated_at timestamp with time zone default now()
        )
        '''
    )


def _annotation_payload(row):
    if not row:
        return {"planned_hours": None, "mapped_ksbs": None, "updated_by": None, "updated_at": None}
    component_id, planned_hours, mapped_ksbs, updated_by, updated_at = row
    return {
        "component_id": component_id,
        "planned_hours": float(planned_hours) if planned_hours is not None else None,
        "mapped_ksbs": mapped_ksbs,
        "updated_by": updated_by,
        "updated_at": updated_at.isoformat() if updated_at else None,
    }


@require_GET
def activity_annotation(request: HttpRequest) -> JsonResponse:
    """Return the auditor-entered KSBs / planned hours for one activity."""
    component = request.GET.get("component", "").strip()
    if not component:
        return JsonResponse({"error": "component is required"}, status=400)
    try:
        with connections[resolve(CONN)].cursor() as cur:
            _ensure_annotation_table(cur)
            cur.execute(
                '''
                select component_id, planned_hours, mapped_ksbs, updated_by, updated_at
                from "Audit".activity_annotations where component_id = %s
                ''',
                [component],
            )
            row = cur.fetchone()
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not read activity annotations.", "details": str(error)}, status=503)
    payload = _annotation_payload(row)
    payload.setdefault("component_id", component)
    return JsonResponse(payload)


@csrf_exempt
def save_activity_annotation(request: HttpRequest) -> JsonResponse:
    """Create/update the auditor-entered KSBs / planned hours for one activity."""
    if request.method != "POST":
        return JsonResponse({"error": "Method not allowed."}, status=405)
    try:
        body = json.loads(request.body or b"{}")
    except ValueError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    component = str(body.get("component_id") or "").strip()
    if not component:
        return JsonResponse({"error": "component_id is required"}, status=400)

    planned_raw = body.get("planned_hours")
    if planned_raw in (None, ""):
        planned_hours = None
    else:
        try:
            planned_hours = float(planned_raw)
        except (TypeError, ValueError):
            return JsonResponse({"error": "planned_hours must be a number"}, status=400)
        if planned_hours < 0 or planned_hours > 100000:
            return JsonResponse({"error": "planned_hours is out of range"}, status=400)

    mapped_ksbs = body.get("mapped_ksbs")
    if mapped_ksbs is not None:
        mapped_ksbs = str(mapped_ksbs).strip()[:5000] or None
    updated_by = str(body.get("updated_by") or "").strip()[:200] or None

    try:
        with connections[resolve(CONN)].cursor() as cur:
            _ensure_annotation_table(cur)
            cur.execute(
                '''
                insert into "Audit".activity_annotations
                    (component_id, planned_hours, mapped_ksbs, updated_by, updated_at)
                values (%s, %s, %s, %s, now())
                on conflict (component_id) do update set
                    planned_hours = excluded.planned_hours,
                    mapped_ksbs = excluded.mapped_ksbs,
                    updated_by = excluded.updated_by,
                    updated_at = now()
                returning component_id, planned_hours, mapped_ksbs, updated_by, updated_at
                ''',
                [component, planned_hours, mapped_ksbs, updated_by],
            )
            row = cur.fetchone()
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not save activity annotation.", "details": str(error)}, status=503)
    return JsonResponse(_annotation_payload(row))


# --- audit-copy activity create/delete overlay -----------------------------
#
# The deployed evidence service owns its source rows and exposes updates, but it
# deliberately has no create/delete API.  New auditor rows and deletions are
# therefore kept as an overlay in the enrolment database.  Deletes are soft
# deletes (tombstones), which preserves the original evidence and gives the UI
# enough information to reverse the corresponding planned/actual totals.

_OVERLAY_CATEGORIES = {"attendance", "assignment", "video", "audio", "reading+quiz"}


def _ensure_activity_overlay_table(cur):
    cur.execute(
        '''
        create table if not exists "Audit".activity_overrides (
            aptem_id bigint not null,
            activity_id text not null,
            operation text not null check (operation in ('created', 'deleted', 'replaced')),
            payload jsonb not null,
            source_payload jsonb,
            updated_by text,
            created_at timestamp with time zone not null default now(),
            updated_at timestamp with time zone not null default now(),
            primary key (aptem_id, activity_id)
        )
        '''
    )
    cur.execute('''alter table "Audit".activity_overrides add column if not exists source_payload jsonb''')
    # Upgrade an early overlay table whose operation check pre-dated date
    # replacement. The DO block is a no-op after the constraint is current.
    cur.execute(
        '''
        do $$
        begin
            if exists (
                select 1 from pg_constraint
                where conname = 'activity_overrides_operation_check'
                  and pg_get_constraintdef(oid) not like '%replaced%'
            ) then
                alter table "Audit".activity_overrides drop constraint activity_overrides_operation_check;
                alter table "Audit".activity_overrides add constraint activity_overrides_operation_check
                    check (operation in ('created', 'deleted', 'replaced'));
            end if;
        end $$
        '''
    )


def _overlay_number(value, field):
    if value in (None, ""):
        return 0.0
    try:
        number = float(value)
    except (TypeError, ValueError):
        raise ValueError(f"{field} must be a number")
    if number < 0 or number > 50:
        raise ValueError(f"{field} must be between 0 and 50 hours")
    return round(number, 4)


def _overlay_timestamp(value, field):
    if value in (None, ""):
        return None
    try:
        parsed = datetime.datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        raise ValueError(f"{field} must be an ISO timestamp")
    return parsed.isoformat()


def _validate_overlay_activity(raw, *, aptem_id, learner_name, activity_id):
    if not isinstance(raw, dict):
        raise ValueError("activity must be an object")
    date = str(raw.get("date") or "").strip()
    try:
        parsed_date = datetime.date.fromisoformat(date)
    except ValueError:
        raise ValueError("date must use YYYY-MM-DD format")
    category = str(raw.get("category") or "").strip().lower()
    if category not in _OVERLAY_CATEGORIES:
        raise ValueError(f"category must be one of: {', '.join(sorted(_OVERLAY_CATEGORIES))}")
    name = str(raw.get("activity") or "").strip()
    if not name:
        raise ValueError("activity is required")
    if len(name) > 500:
        raise ValueError("activity must be at most 500 characters")

    planned = _overlay_number(raw.get("planned"), "planned")
    actual = _overlay_number(raw.get("actual"), "actual")
    started = _overlay_timestamp(raw.get("timestamp_from"), "timestamp_from")
    completed = _overlay_timestamp(raw.get("timestamp_to"), "timestamp_to")
    if started and completed:
        start_value = datetime.datetime.fromisoformat(started)
        end_value = datetime.datetime.fromisoformat(completed)
        if start_value.tzinfo is None and end_value.tzinfo is not None or start_value.tzinfo is not None and end_value.tzinfo is None:
            raise ValueError("timestamps must use the same timezone style")
        if end_value < start_value:
            raise ValueError("timestamp_to must be after timestamp_from")

    display = str(raw.get("timestamp_display") or "").strip()[:100]
    if not display:
        if started and completed:
            display = f"{datetime.datetime.fromisoformat(started):%H:%M}–{datetime.datetime.fromisoformat(completed):%H:%M}"
        elif actual > 0:
            display = "input"

    return {
        "activity_id": activity_id,
        "learner_id": aptem_id,
        "learner_name": learner_name,
        "date": parsed_date.isoformat(),
        "month": parsed_date.strftime("%Y-%m"),
        "month_label": parsed_date.strftime("%B %Y"),
        "category": category,
        "activity": name,
        "activity_subtitle": str(raw.get("activity_subtitle") or "").strip()[:2000] or None,
        "planned": planned,
        "actual": actual,
        "timestamp_from": started,
        "timestamp_to": completed,
        "timestamp_display": display,
        "completed": bool(raw.get("completed", actual > 0)),
        "ksbs": raw.get("ksbs") if isinstance(raw.get("ksbs"), dict) else {"K": [], "S": [], "B": []},
        "iframe_url": None,
        "not_accepted": bool(raw.get("not_accepted", False)),
        "reporting_week_label": str(raw.get("reporting_week_label") or "").strip()[:200] or None,
        "audit_created": bool(raw.get("audit_created", str(activity_id).startswith("audit:"))),
    }


def _overlay_learner(aptem_id):
    """Resolve a writable audit learner without the legacy PCP-only filter.

    The journal reads the expanded live multi-programme cohort, while
    ``_load_rows`` intentionally remains restricted to the original Project
    Controls report.  Overlay writes therefore validate directly against the
    shared ``Audit.learner_match`` source by Aptem ID.  This keeps the legacy
    report endpoints unchanged while allowing create/date-replace/delete for
    every learner that the expanded audit workspace can surface.
    """
    with connections[resolve(CONN)].cursor() as cur:
        cur.execute(
            '''
            select learner_name
            from "Audit".learner_match
            where aptem_id = %s
            limit 1
            ''',
            [aptem_id],
        )
        row = cur.fetchone()
    if not row:
        return None
    return {"aptem_id": aptem_id, "name": row[0] or f"Learner {aptem_id}"}


@csrf_exempt
def activity_overrides(request: HttpRequest) -> JsonResponse:
    """List/create/update/soft-delete audit-created activity overlays."""
    if request.method == "GET":
        raw_id = request.GET.get("aptem_id", "").strip()
        try:
            aptem_id = int(raw_id) if raw_id else None
        except ValueError:
            return JsonResponse({"error": "aptem_id must be an integer"}, status=400)
        try:
            with connections[resolve(CONN)].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                if aptem_id is None:
                    cur.execute('''select aptem_id, activity_id, operation, payload, source_payload, updated_by, updated_at from "Audit".activity_overrides order by updated_at''')
                else:
                    cur.execute('''select aptem_id, activity_id, operation, payload, source_payload, updated_by, updated_at from "Audit".activity_overrides where aptem_id = %s order by updated_at''', [aptem_id])
                rows = cur.fetchall()
        except (KeyError, DatabaseError) as error:
            return JsonResponse({"error": "Could not read activity overrides.", "details": str(error)}, status=503)
        return JsonResponse({"items": [
            {"aptem_id": row[0], "activity_id": row[1], "operation": row[2], "payload": row[3], "source_payload": row[4], "updated_by": row[5], "updated_at": row[6].isoformat() if row[6] else None}
            for row in rows
        ]})

    if request.method not in {"POST", "PUT", "PATCH", "DELETE"}:
        return JsonResponse({"error": "Method not allowed."}, status=405)
    try:
        body = json.loads(request.body or b"{}")
        aptem_id = int(body.get("aptem_id"))
    except (TypeError, ValueError):
        return JsonResponse({"error": "aptem_id must be an integer"}, status=400)
    try:
        learner = _overlay_learner(aptem_id)
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not validate learner.", "details": str(error)}, status=503)
    if not learner:
        return JsonResponse({"error": "Learner is outside the audit-copy cohort."}, status=404)

    updated_by = str(body.get("updated_by") or "").strip()[:200] or None
    activity_id = str(body.get("activity_id") or "").strip()
    source_payload = None
    try:
        if request.method == "POST":
            activity_id = f"audit:{uuid.uuid4()}"
            payload = _validate_overlay_activity(
                body.get("activity"), aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
            )
            operation = "created"
        elif request.method == "PUT":
            if not activity_id:
                raise ValueError("activity_id is required")
            with connections[resolve(CONN)].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                cur.execute('''select source_payload from "Audit".activity_overrides where aptem_id = %s and activity_id = %s and operation = 'replaced' ''', [aptem_id, activity_id])
                existing = cur.fetchone()
            raw_source = existing[0] if existing and existing[0] else body.get("snapshot")
            if not isinstance(raw_source, dict):
                raise ValueError("snapshot is required for a reversible date change")
            source_payload = _validate_overlay_activity(
                raw_source, aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
            )
            payload = _validate_overlay_activity(
                body.get("activity"), aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
            )
            payload["audit_replaced"] = True
            operation = "replaced"
        elif request.method == "PATCH":
            if not activity_id.startswith("audit:"):
                return JsonResponse({"error": "Only audit-created activities can be patched here."}, status=400)
            with connections[resolve(CONN)].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                cur.execute('''select payload from "Audit".activity_overrides where aptem_id = %s and activity_id = %s and operation = 'created' ''', [aptem_id, activity_id])
                existing = cur.fetchone()
            if not existing:
                return JsonResponse({"error": "Audit activity was not found."}, status=404)
            merged = {**(existing[0] or {}), **(body.get("patch") or {})}
            payload = _validate_overlay_activity(
                merged, aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
            )
            operation = "created"
        else:
            if not activity_id:
                raise ValueError("activity_id is required")
            raw_snapshot = body.get("snapshot")
            with connections[resolve(CONN)].cursor() as cur:
                _ensure_activity_overlay_table(cur)
                cur.execute('''select source_payload from "Audit".activity_overrides where aptem_id = %s and activity_id = %s and operation = 'replaced' ''', [aptem_id, activity_id])
                existing = cur.fetchone()
            if existing and existing[0]:
                raw_snapshot = existing[0]
            if not isinstance(raw_snapshot, dict):
                raise ValueError("snapshot is required for a reversible deletion")
            payload = _validate_overlay_activity(
                raw_snapshot, aptem_id=aptem_id,
                learner_name=learner["name"], activity_id=activity_id,
            )
            operation = "deleted"
    except ValueError as error:
        return JsonResponse({"error": str(error)}, status=400)

    try:
        with connections[resolve(CONN)].cursor() as cur:
            _ensure_activity_overlay_table(cur)
            cur.execute(
                '''
                insert into "Audit".activity_overrides (aptem_id, activity_id, operation, payload, source_payload, updated_by)
                values (%s, %s, %s, %s::jsonb, %s::jsonb, %s)
                on conflict (aptem_id, activity_id) do update set
                    operation = excluded.operation, payload = excluded.payload,
                    source_payload = excluded.source_payload,
                    updated_by = excluded.updated_by, updated_at = now()
                returning updated_at
                ''',
                [aptem_id, activity_id, operation, json.dumps(payload), json.dumps(source_payload) if source_payload else None, updated_by],
            )
            updated_at = cur.fetchone()[0]
    except (KeyError, DatabaseError) as error:
        return JsonResponse({"error": "Could not save activity override.", "details": str(error)}, status=503)
    return JsonResponse({
        "ok": True, "aptem_id": aptem_id, "activity_id": activity_id,
        "operation": operation, "payload": payload,
        "updated_by": updated_by, "updated_at": updated_at.isoformat(),
    })
